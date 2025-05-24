import re
from typing import Dict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement # For shading
from docx.oxml.ns import qn # For shading

def fill_word_document_from_llm_data(template_path: str, data: Dict[str, str], output_path: str) -> None:
    """
    Fills the placeholders in the Word document with data (typically from LLM).
    Handles checkbox placeholders (ending in '_check') and simple text placeholders.
    Applies basic formatting (bold, black font, yellow background) to changed cells.

    Args:
        template_path: Path to the .docx template.
        data: Dictionary of placeholder keys to their string values ("YES"/"NO" for checks).
        output_path: Path to save the filled document.
    """
    try:
        doc = Document(template_path)
        checked_symbol = "☒"  # Unicode for checked box
        unchecked_symbol = "☐"  # Unicode for unchecked box

        # Placeholders are already flat (e.g., from LLM response)
        for i, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    if '{{' not in cell.text: # Quick check
                        continue
                    
                    original_cell_text = cell.text
                    modified_cell_text = original_cell_text
                    cell_was_modified_to_checked = False # Track if a checkbox was set to YES

                    for key, value_str in data.items():
                        # Regex to find placeholder, allowing for whitespace within {{...}}
                        placeholder_regex_str = r"{{\s*" + re.escape(key) + r"\s*}}"
                        placeholder_regex = re.compile(placeholder_regex_str)
                        
                        if placeholder_regex.search(modified_cell_text):
                            replacement = ""
                            is_check_field = key.endswith("_check")
                            is_checked_yes = is_check_field and value_str.upper() == "YES"

                            if is_check_field:
                                replacement = checked_symbol if is_checked_yes else unchecked_symbol
                            else:
                                replacement = str(value_str) # For non-checkbox placeholders
                            
                            new_text, num_subs = placeholder_regex.subn(replacement, modified_cell_text)
                            if num_subs > 0:
                                modified_cell_text = new_text
                                if is_checked_yes: # Only flag for highlight if it became a checked box
                                    cell_was_modified_to_checked = True 
                                elif not is_check_field and value_str: # Also consider non-empty text fields as modified for formatting
                                     cell_was_modified_to_checked = True # Re-using flag, but logic is for any content change
                    
                    if original_cell_text != modified_cell_text: # If any text change occurred
                        cell.text = modified_cell_text 
                        # Apply text formatting to all modified cells
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)  
                                run.bold = True
                                run.font.size = Pt(12)
                        
                        # Apply yellow background ONLY if a checkbox was set to YES or a text field was filled
                        if cell_was_modified_to_checked:
                            tcPr = cell._element.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:fill'), 'FFFF00') 
                            tcPr.append(shd)

        doc.save(output_path)
        print(f"Successfully created filled document: {output_path}")

    except Exception as e:
        print(f"Error in fill_word_document_from_llm_data: {e}")
        raise # Re-raise exception to be caught by main

# Example Usage (for testing this module directly)
if __name__ == '__main__':
    import os
    mock_data_for_filler = {
        "plc_b&r_check": "YES",
        "hmi_size10_check": "YES",
        "hmi_size15_check": "NO",
        "vd_f_check": "YES",
        "some_other_text_field": "Example Text Value"
    }
    test_template = 'template.docx'
    test_output = 'filler_test_output.docx'

    if os.path.exists(test_template):
        print(f"Testing document filler with template: {test_template}")
        fill_word_document_from_llm_data(test_template, mock_data_for_filler, test_output)
    else:
        print(f"Test template '{test_template}' not found.") 