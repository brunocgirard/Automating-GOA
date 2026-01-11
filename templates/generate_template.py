import os
import sys
import re
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt

# Ensure we can import from src
sys.path.append(os.getcwd())

try:
    from src.utils.template_utils import DEFAULT_EXPLICIT_MAPPINGS
except ImportError:
    print("Could not import DEFAULT_EXPLICIT_MAPPINGS. Make sure you are running this from the project root.")
    sys.exit(1)

OUTLINE_PATH = "full_fields_outline.md"
TEMPLATE_PATH = os.path.join("templates", "template.docx")
OUTPUT_PATH = os.path.join("templates", "comprehensive_template.docx")

def normalize_key(text: str) -> str:
    """Converts text to a snake_case key."""
    # Remove parenthetical info like (checkbox), (text)
    text = re.sub(r'\s*\(.*?\)', '', text)
    text = text.lower().strip()
    text = re.sub(r'[^a-z0-9]+', '_', text)
    return text.strip('_')

def determine_type_suffix(line: str) -> str:
    """Determines the suffix based on (type) in the line."""
    lower_line = line.lower()
    if "(checkbox)" in lower_line:
        return "_check"
    elif "(qty)" in lower_line:
        return "_qty"
    elif "(text)" in lower_line or "(text area)" in lower_line:
        return "_text"
    return "_text" # Default

def parse_outline(file_path: str) -> List[Dict]:
    """Parses the markdown outline into a structured list of fields."""
    fields = []
    current_section = "General"
    # Stack to track hierarchy: [(indent_level, text)]
    # Level 0 is the section (handled separately), so stack starts at level 1 items
    stack = [] 

    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip()
            if not line:
                continue
            
            # Section Header
            if line.startswith("## "):
                current_section = line[3:].strip()
                stack = [] # Reset stack on new section
                continue
            
            # List Item
            match = re.match(r'^(\s*)-\s+(.*)', line)
            if match:
                indent_str = match.group(1)
                content = match.group(2)
                
                # Calculate indent level (assuming 2 spaces per level)
                indent_level = len(indent_str) // 2
                
                # Adjust stack
                while stack and stack[-1][0] >= indent_level:
                    stack.pop()
                
                # Determine if this item is a "group" or a "field"
                # Heuristic: If it has children, it's a group. 
                # But we process line by line, so we don't know yet.
                # However, for the purpose of generating keys, we can treat everything as potential context.
                
                parent_context = [item[1] for item in stack]
                stack.append((indent_level, content))
                
                # Construct description for mapping lookup
                # Clean content for description (remove type info)
                clean_content = re.sub(r'\s*\(.*?\)', '', content).strip()
                
                # Build hierarchy string: Section - [Parent -] Content
                hierarchy = [current_section] + [re.sub(r'\s*\(.*?\)', '', p).strip() for p in parent_context] + [clean_content]
                description = " - ".join(hierarchy)
                
                # Determine type
                suffix = determine_type_suffix(content)
                
                fields.append({
                    "section": current_section,
                    "hierarchy": hierarchy,
                    "description": description,
                    "raw_content": content,
                    "suffix": suffix,
                    "indent": indent_level
                })

    return fields

def find_existing_key(description: str, mappings: Dict[str, str]) -> Optional[str]:
    """Looks up the description in the existing mappings."""
    # Create a reverse mapping for easier lookup
    # Normalize descriptions for comparison (lowercase, remove special chars)
    
    def clean_desc(d):
        return re.sub(r'[^a-z0-9]', '', d.lower())
    
    target = clean_desc(description)
    
    for key, map_desc in mappings.items():
        if clean_desc(map_desc) == target:
            return key
            
    return None

def generate_key(field_data: Dict) -> str:
    """Generates a new key based on hierarchy."""
    # Use the last few parts of the hierarchy to keep it somewhat short but unique
    # e.g. section_parent_item_suffix
    
    parts = [normalize_key(p) for p in field_data['hierarchy']]
    
    # If section is long, maybe abbreviate? For now, keep it full to ensure uniqueness.
    key_base = "_".join(parts)
    
    # Ensure it doesn't end with the suffix if it's already there (e.g. if content was "Check (checkbox)")
    suffix = field_data['suffix']
    if key_base.endswith(suffix):
        return key_base
    
    return f"{key_base}{suffix}"

def main():
    print(f"Parsing outline from {OUTLINE_PATH}...")
    outline_fields = parse_outline(OUTLINE_PATH)
    print(f"Found {len(outline_fields)} items in outline.")
    
    print(f"Loading template from {TEMPLATE_PATH}...")
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Error: Template file not found at {TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)
    
    # Extract existing placeholders to avoid duplicates
    existing_placeholders = set()
    regex = re.compile(r"{{\s*(.*?)\s*}}")
    for p in doc.paragraphs:
        existing_placeholders.update(regex.findall(p.text))
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                existing_placeholders.update(regex.findall(c.text))
                
    print(f"Found {len(existing_placeholders)} existing placeholders in template.")
    
    # Process fields
    fields_to_add = []
    mapped_count = 0
    new_key_count = 0
    
    # Create a map of generated keys to ensure uniqueness within this run
    generated_keys = set()

    for field in outline_fields:
        # 1. Try to find existing key
        key = find_existing_key(field['description'], DEFAULT_EXPLICIT_MAPPINGS)
        
        if key:
            mapped_count += 1
        else:
            # 2. Generate new key
            key = generate_key(field)
            new_key_count += 1
        
        # Ensure uniqueness if we have collisions in generation
        original_key = key
        counter = 1
        while key in generated_keys:
            key = f"{original_key}_{counter}"
            counter += 1
        generated_keys.add(key)
        
        # 3. Check if in document
        if key not in existing_placeholders:
            fields_to_add.append((key, field))

    print(f"Mapped {mapped_count} fields to existing keys.")
    print(f"Generated {new_key_count} new keys.")
    print(f"Identified {len(fields_to_add)} fields missing from the template.")
    
    if fields_to_add:
        doc.add_page_break()
        doc.add_heading('Supplemental Fields', level=1)
        doc.add_paragraph("The following fields were defined in the outline but missing from the original template.")
        
        current_section = ""
        for key, field in fields_to_add:
            if field['section'] != current_section:
                doc.add_heading(field['section'], level=2)
                current_section = field['section']
            
            # Add the field
            # Format: Label: {{key}}
            label = field['raw_content']
            # Indent based on level
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(field['indent'] * 12) # 12pt per level
            
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            p.add_run(f"{{{{{key}}}}}")
            
    print(f"Saving new template to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print("Done.")

if __name__ == "__main__":
    main()
