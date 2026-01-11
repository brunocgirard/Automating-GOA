import re
from typing import Dict

from docx import Document
from docx.shared import Pt, RGBColor


def fill_word_document_from_llm_data(template_path: str, data: Dict[str, str], output_path: str) -> None:
    """
    Fills placeholders in a Word document with provided data.
    Checkbox placeholders end with "_check" and are rendered as checked/unchecked symbols.
    Modified cells are bolded in black without background highlight.
    """
    try:
        doc = Document(template_path)
        checked_symbol = "☑"  # Unicode ballot box with check (U+2611)
        unchecked_symbol = "☐"  # Unicode ballot box (U+2610)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{" not in cell.text:
                        continue

                    original_cell_text = cell.text
                    modified_cell_text = original_cell_text

                    for key, value_str in data.items():
                        placeholder_regex = re.compile(r"{{\s*" + re.escape(key) + r"\s*}}")
                        if not placeholder_regex.search(modified_cell_text):
                            continue

                        is_check_field = key.endswith("_check")
                        is_checked_yes = is_check_field and str(value_str).upper() == "YES"

                        if is_check_field:
                            replacement = checked_symbol if is_checked_yes else unchecked_symbol
                        else:
                            replacement = str(value_str)

                        modified_cell_text, _ = placeholder_regex.subn(replacement, modified_cell_text)

                    if original_cell_text != modified_cell_text:
                        cell.text = modified_cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.bold = True
                                run.font.size = Pt(12)

        doc.save(output_path)
        print(f"Successfully created filled document: {output_path}")

    except Exception as e:
        print(f"Error in fill_word_document_from_llm_data: {e}")
        raise


if __name__ == '__main__':
    import os

    mock_data_for_filler = {
        "plc_b&r_check": "YES",
        "hmi_size10_check": "YES",
        "hmi_size15_check": "NO",
        "vd_f_check": "YES",
        "some_other_text_field": "Example Text Value",
    }
    test_template = 'template.docx'
    test_output = 'filler_test_output.docx'

    if os.path.exists(test_template):
        print(f"Testing document filler with template: {test_template}")
        fill_word_document_from_llm_data(test_template, mock_data_for_filler, test_output)
    else:
        print(f"Test template '{test_template}' not found.")
