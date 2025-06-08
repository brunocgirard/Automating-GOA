import re
from typing import List, Set, Dict, Optional
from docx import Document
import os # Added for __main__ example

# Define explicit placeholder mappings for the default template
DEFAULT_EXPLICIT_MAPPINGS = {
    # Euro Guarding specific mappings
    "eg_none_check": "Euro Guarding - None",
    "eg_pmtg_check": "Euro Guarding - Panel material - Tempered glass",
    "eg_pnl_check": "Euro Guarding - Panel material - Lexan",
    "eg_stkw_check": "Euro Guarding - Switch type - Key switch",
    "eg_stm_check": "Euro Guarding - Switch type - Magnetic",
    "eg_sto_check": "Euro Guarding - Switch type - Other",
    "eg_tcnone_check": "Euro Guarding - Top cover - None",
    "eg_tcyes_check": "Euro Guarding - Top cover - Yes",
    "eg_rc_check": "Euro Guarding - Reject cover",
    "eg_rcnone_check": "Euro Guarding - Reject cover - None",
    
    # Coding and Inspection System Specifications mappings
    "ci_chs_check": "Coding and Inspection System Specifications - Coder - hot stamp",
    "ci_ci15_check": "Coding and Inspection System Specifications - Videojet - ink jet vj 1520",
    "ci_cl_check": "Coding and Inspection System Specifications - Coder - laser",
    "ci_cnone_check": "Coding and Inspection System Specifications - Coder - none",
    "ci_ctj_check": "Coding and Inspection System Specifications - Coder - thermal ink jet",
    "ci_lfe_check": "Coding and Inspection System Specifications - Laser - fume extraction",
    "ci_lfm": "Coding and Inspection System Specifications - Laser - fume extraction model",
    "ci_lo_check": "Coding and Inspection System Specifications - Laser - other",
    "ci_lv_check": "Coding and Inspection System Specifications - Laser - videojet",
    "ci_lvm": "Coding and Inspection System Specifications - Laser - videojet model",
    "ci_p2d_check": "Coding and Inspection System Specifications - Print - 2d",
    "ci_pbc_check": "Coding and Inspection System Specifications - Print - bar code",
    "ci_pep_check": "Coding and Inspection System Specifications - Print - exp. date",
    "ci_pl_check": "Coding and Inspection System Specifications - Print - lot",
    "ci_pnone_check": "Coding and Inspection System Specifications - Print - none",
    "ci_po_check": "Coding and Inspection System Specifications - Print - other",
    "ci_ppo_check": "Coding and Inspection System Specifications - Print - package on",
    "ci_v33_check": "Coding and Inspection System Specifications - Videojet - laser 3330",
    "ci_v85_check": "Coding and Inspection System Specifications - Videojet - 8520",
    "ci_vb_check": "Coding and Inspection System Specifications - Vision - barcode",
    "ci_vc_check": "Coding and Inspection System Specifications - Vision - cognex",
    "ci_vcom": "Coding and Inspection System Specifications - Vision - comment",
    "ci_vd651_check": "Coding and Inspection System Specifications - Videojet - dataflex 6530 (107wx75l)",
    "ci_vd65_check": "Coding and Inspection System Specifications - Videojet -dataflex 6530 (53wx75l)",
    "ci_vdxsx": "Coding and Inspection System Specifications - Videojet - for dx (rh) for sx (lh)",
    "ci_vnone_check": "Coding and Inspection System Specifications - Vision - none",
    "ci_vo_check": "Coding and Inspection System Specifications - Vision - other",
    "ci_vocr_check": "Coding and Inspection System Specifications - Vision - ocr",
    "ci_vocv_check": "Coding and Inspection System Specifications - Vision - ocv",
    "ci_vos_check": "Coding and Inspection System Specifications - Vision - orientation/ skew",
    "ci_vow_check": "Coding and Inspection System Specifications - Vision - on web",
    "ci_vp_check": "Coding and Inspection System Specifications - Vision - presence",
    "ci_vqty": "Coding and Inspection System Specifications - Vision - qty",
    "ci_vpo_check": "Coding and Inspection System Specifications - Vision - position/ placement",
    "ci_vtq_check": "Coding and Inspection System Specifications - Vision - torque",
    "ci_ctt_check": "Coding and Inspection System Specifications - Coder - thermal transfer",
    "ci_vob_check": "Coding and Inspection System Specifications - Vision - on bottle",
    
    # Warranty & Install & Spares mappings
    "sk_1yr_check": "Warranty & Install & Spares - Spare kit - 1yr",
    "sk_2yr_check": "Warranty & Install & Spares - Spare kit - 2yr",
    "sk_none_check": "Warranty & Install & Spares - Spare kit - none",
    "stpc_none_check": "Warranty & Install & Spares - Start-up commissioning - none",
    "stpc_yes": "Warranty & Install & Spares - Start-up commissioning - yes – # of days",
    "wi_war1_check": "Warranty & Install & Spares - Warranty - 1yr",
    "wi_war2_check": "Warranty & Install & Spares - Warranty - 2yr",
    "wrts_y_check": "Warranty & Install & Spares - Remote tech service - yes",
    "wrts_none_check": "Warranty & Install & Spares - Remote tech service - none",
    
    # Packaging & Transport mappings
    "pt_pbc_check": "Packaging & Transport - Packaging by - capmatic",
    "pt_pbcu_check": "Packaging & Transport - Packaging by - customer / not incl.",
    "pt_ptc_check": "Packaging & Transport - Packaging type - crate",
    # "pt_pts_check": "Packaging & Transport - Packaging type - skid",
    "pt_ptsw_check": "Packaging & Transport - Packaging type -  sea worthy",
    "pt_tcc_check": "Packaging & Transport - Transport charges - capmatic",
    "pt_tccu_check": "Packaging & Transport - Transport charges - customer",
    
    # Validation Documents mappings
    "vd_d_check": "Validation Documents - Design Qualification (DQ) Documentation",
    "vd_f_check": "Validation Documents - Factory Acceptance Test (FAT) Protocol Package - Includes documentation, support, and customer review",
    "vd_fd_check": "Validation Documents - Functional/Design Specification (FS/DS)",
    "vd_h_check": "Validation Documents - Hardware/Software Design Specification (HDS/SDS)",
    "vd_i_check": "Validation Documents - Installation/Operational Qualification (IQ/OQ)",
    "vd_s_check": "Validation Documents - Site Acceptance Test (SAT) Protocol Package - Includes documentation, support, and customer review",
    
    # Labeling System Specifications mappings
    "ls_a3p_check": "Labeling System Specifications - Application sys - 3-panel",
    "ls_a4p_check": "Labeling System Specifications - Application sys - 4-panel",
    "ls_a5p_check": "Labeling System Specifications - Application sys - 5-panel",
    "ls_abqty": "Labeling System Specifications - Application sys - belts qty",
    "ls_acl_check": "Labeling System Specifications - Application sys - case label",
    "ls_ap_check": "Labeling System Specifications - Application sys - prism",
    "ls_atl_check": "Labeling System Specifications - Arm type - \"l\" shaped",
    "ls_ats_check": "Labeling System Specifications - Arm type - standard",
    "ls_awa_check": "Labeling System Specifications - Application sys - wrap around",
    "ls_awb_check": "Labeling System Specifications - Application sys - wipe front & back",
    "ls_awp_check": "Labeling System Specifications - Application sys - wrap around in puck",
    "ls_ee_check": "Labeling System Specifications - eagle eye plus sensor (clr label)",
    "ls_lhdx_check": "Labeling System Specifications - Label head orientation - dx right side",
    "ls_lhl100_check": "Labeling System Specifications - Label hd model - ls100",
    "ls_lhl200_check": "Labeling System Specifications - Label hd model - ls200",
    "ls_lhsx_check": "Labeling System Specifications - Label head orientation - sx left side",
    "ls_ll_check": "Labeling System Specifications - low label sensor",
    "ls_none_check": "Labeling System Specifications - none",
    "ls_sd300_check": "Labeling System Specifications - Support reel dia - 300 mm",
    "ls_sd380_check": "Labeling System Specifications - Support reel dia - 380 mm",
    "ls_swb_check": "Labeling System Specifications - Separator wheel - belt",
    "ls_swnone_check": "Labeling System Specifications - Separator wheel - none",
    "ls_sws_check": "Labeling System Specifications - Separator wheel - starwheel",
    "ls_sww_check": "Labeling System Specifications - Separator wheel - wheel",
    "ls_thd_check": "Labeling System Specifications - top hold down",
    "ls_thdba_check": "Labeling System Specifications - bottle aligner",
    "ls_thdtf_check": "Labeling System Specifications - twin feedscrew",
    "ls_tl_check": "Labeling System Specifications - Type - linear (inline)",
    "ls_tr_check": "Labeling System Specifications - Type - rotary",
    
    # BeltStar System Specifications mappings
    "bs_cpac_check": "BeltStar System Specifications - Cap Placement - AC",
    "bs_cpbo_check": "BeltStar System Specifications - Cap Placement - Belts – on the fly",
    "bs_cpnone_check": "BeltStar System Specifications - Cap Placement - None",
    "bs_csds_check": "BeltStar System Specifications - Cap Sorting - Docking Station",
    "bs_cse_check": "BeltStar System Specifications - Cap Sorting - Elevator",
    "bs_manone_check": "BeltStar System Specifications - Motorized Adj. - None",
    "bs_may_check": "BeltStar System Specifications - Motorized Adj. - Yes",
    "bs_none_check": "BeltStar System Specifications - None",
    "bs_tadb_check": "BeltStar System Specifications - Torque - AC motor DC Brake",
    "bs_tam_check": "BeltStar System Specifications - Torque - Air motor",
    "bs_tb_check": "BeltStar System Specifications - Torque - Belts",
    "bs_tf_check": "BeltStar System Specifications - Torque - Feedback",
    "bs_tht_check": "BeltStar System Specifications - Torque - HMI Adj. Torque",
    "bs_tmc_check": "BeltStar System Specifications - Torque - Magnet clutch",
    "bs_ts_check": "BeltStar System Specifications - Torque - Servo",
    "bs_tcb": "BeltStar System Specifications - Torque - Copy of bottle",
    "bs_tcn": "BeltStar System Specifications - Torque - copy of neck",
    "bs_ttr": "BeltStar System Specifications - Torque - Torque Range (in-lbs)",
    
    # Capping System Specifications mappings
    "cs_cdc_check": "Capping System Specifications - Centering device - cone",
    "cs_cdn_check": "Capping System Specifications - Centering device - neck",
    "cs_cdnone_check": "Capping System Specifications - Centering device - none",
    "cs_cdt_check": "Capping System Specifications - Centering device - tube",
    "cs_cpof_check": "Capping System Specifications - Cap placement - on the fly",
    "cs_cppp_check": "Capping System Specifications - Cap placement - pick & place",
    "cs_cppt_check": "Capping System Specifications - Cap placement - push through",
    "cs_cpqty": "Capping System Specifications - Cap placement - qty",
    "cs_cpr_check": "Capping System Specifications - Cap placement - rotary",
    "cs_cpra_check": "Capping System Specifications - Cap placement - Rotation - air",
    "cs_cprs_check": "Capping System Specifications - Cap placement - Rotation - servo",
    "cs_cpsd_check": "Capping System Specifications - Cap placement - servo up down",
    "cs_cpsi_check": "Capping System Specifications - Cap placement - servo index",
    "cs_csac_check": "Capping System Specifications - Cap sorting - acoustic cover",
    "cs_csc_check": "Capping System Specifications - Cap sorting - centrifugal",
    "cs_csds_check": "Capping System Specifications - Cap sorting - docking station",
    "cs_csm_check": "Capping System Specifications - Cap Sorting - Mechanical",
    "cs_csvb_check": "Capping System Specifications - Cap sorting - vibratory bowl (xmm)",
    "cs_none_check": "Capping System Specifications - none",
    "cs_tam_check": "Capping System Specifications - Torque - air motor",
    "cs_tat_check": "Capping System Specifications - Torque - appl. torque system tool (30 or 44)",
    "cs_tcn": "Capping System Specifications - Torque - copy of neck",
    "cs_tf_check": "Capping System Specifications - Torque - feedback",
    "cs_tmc_check": "Capping System Specifications - Torque - magnet clutch",
    "cs_tqty": "Capping System Specifications - Torque - qty",
    "cs_ts_check": "Capping System Specifications - Torque - servo",
    "cs_tsc_check": "Capping System Specifications - Torque - servo calibration sys.",
    "cs_ttr": "Capping System Specifications - Torque - torque range (in-lbs)",
    "cs_bfc": "Capping System Specifications - Bulk feeder - capacity (ft3)",
    "cs_bfc_check": "Capping System Specifications - Bulk feeder - cover",
    "cs_bfeg_check": "Capping System Specifications- Bulk feeder - elevator (giraffe)",
    "cs_bfnone_check": "Capping System Specifications- Bulk feeder - none",
    "cs_bfpv_check": "Capping System Specifications- Bulk feeder - pedestal vibratory",
    "cs_cse_check": "Capping System Specifications - Cap sorting elevator",
    "cs_tac_check": "Capping System Specifications - Tube aligner - cone",
    "cs_tagd_check": "Capping System Specifications - Tube aligner - grip & dive",
    "cs_tanone_check": "Capping System Specifications - Tube aligner - none",
    "cs_tcb": "Capping System Specifications - Torque - copy of bottle",
    
    # Plugging System Specifications mappings
    "ps_none_check": "Plugging System Specifications - none",
    "ps_plm_check": "Plugging System Specifications - Plug Placement - Mechanical",
    "ps_psm_check": "Plugging System Specifications - Plug Sorting - Mechanical",
    "ps_bfc": "Plugging System Specifications - Bulk feeder - capacity (ft3)",
    "ps_bfe_check": "Plugging System Specifications - Bulk feeder - electric",
    "ps_bfeg_check": "Plugging System Specifications - Bulk feeder - elevator (giraffe)",
    "ps_bfnone_check": "Plugging System Specifications - Bulk feeder - none",
    "ps_bfp_check": "Plugging System Specifications - Bulk feeder - pneumatic",
    "ps_bfpv_check": "Plugging System Specifications - Bulk feeder - pedestal vibratory",
    "ps_plpp_check": "Plugging System Specifications - Plug placement - pick & place",
    "ps_plpt_check": "Plugging System Specifications - Plug placement - push through",
    "ps_plqty": "Plugging System Specifications - Plug placement - qty",
    "ps_plr_check": "Plugging System Specifications - Plug placement - rotary",
    "ps_plv_check": "Plugging System Specifications - Plug placement - vacuum",
    "ps_psac_check": "Plugging System Specifications - Plug sorting acoustic cover",
    "ps_psc_check": "Plugging System Specifications - Plug sorting - centrifugal",
    "ps_psds_check": "Plugging System Specifications - Plug sorting - docking station",
    "ps_psnone_check": "Plugging System Specifications - Plug sorting - none",
    "ps_psvb_check": "Plugging System Specifications - Plug sorting - vibratory bowl (?mm)",
    "ps_pse_check": "Plugging System Specifications - Plug sorting - elevator",
    
    # Cottoner mappings
    "c_cn_check": "Cottoner - Cotton Bin - No",
    "c_cy_check": "Cottoner - Cotton Bin - Yes",
    "c_none_check": "Cottoner - None",
    "c_sp_check": "Cottoner - Sensing - Presence",
    "c_sh_check": "Cottoner - Sensing - High",

    # Desiccant mappings
    "d_bf_pv_check": "Desiccant - Bulk feeder - pedestal vibratory",
    "d_bfc_check": "Desiccant - Bulk feeder - cover",
    "d_bfe_check": "Desiccant - Bulk feeder - elevator",
    "d_bfi_check": "Desiccant - Bulk feeder - internal",
    "d_none_check": "Desiccant - none",
    "d_tc_check": "Desiccant - Type - cannister",
    "d_trp_check": "Desiccant - Type - roll / pouch",

    # Gas Purge mappings
    "gp_laf_check": "Gas Purge - Location/Type - after fill",
    "gp_latf_check": "Gas Purge - Location/Type - at fill st.",
    "gp_lbf_check": "Gas Purge - Location/Type - before fill st",
    "gp_lqty": "Gas Purge - Location/Type - qty",
    "gp_lt_check": "Gas Purge - Location/Type - tunnel",
    "gp_none_check": "Gas Purge - none",
    "gp_ta_check": "Gas Purge - Type - argon",
    "gp_tn_check": "Gas Purge - Type - nitrogen",

    # Liquid Filling System Specifications mappings
    "lf_cnone_check": "Liquid Filling System Specifications - Cleaning - none",
    "lf_cns": "Liquid Filling System Specifications - Cleaning - no. of stations",
    "lf_coc_check": "Liquid Filling System Specifications - Cleaning - ionization cleaning",
    "lf_ctg_check": "Liquid Filling System Specifications - Cleaning - touch & go (c.i.p. man. sys)",
    "lf_cvac_check": "Liquid Filling System Specifications - Cleaning - vacuum air cleaning",
    "lf_gsds_check": "Liquid Filling System Specifications - Gutter system - drip style",
    "lf_gsnone_check": "Liquid Filling System Specifications - Gutter system - none",
    "lf_gsp_check": "Liquid Filling System Specifications - Gutter system - prime/purge/rinse",
    "lf_nb1/2_check": "Liquid Filling System Specifications - Nozzle body size - 1/2\"",
    "lf_nb1/4_check": "Liquid Filling System Specifications - Nozzle body size - 1/4\"",
    "lf_nb10_check": "Liquid Filling System Specifications - Nozzle body size - 10mm",
    "lf_nb12_check": "Liquid Filling System Specifications - Nozzle body size - 12mm",
    "lf_nb1_check": "Liquid Filling System Specifications - Nozzle body size - 1\"",
    "lf_nb3/8_check": "Liquid Filling System Specifications - Nozzle body size - 3/8\"",
    "lf_nb5/8_check": "Liquid Filling System Specifications - Nozzle body size - 5/8\"",
    "lf_nb6_check": "Liquid Filling System Specifications - Nozzle body size - 6 mm",
    "lf_nb8_check": "Liquid Filling System Specifications - Nozzle body size - 8mm",
    "lf_nb3/4_check": "Liquid Filling System Specifications - Nozzle body type - 3/4\"",
    "lf_nson_check": "Liquid Filling System Specifications - offset nozzle bar",
    "lf_ntd_check": "Liquid Filling System Specifications - Nozzle type - double wall nitro & fill",
    "lf_ntdq": "Liquid Filling System Specifications - Nozzle type - double wall nitro & fill - qty",
    "lf_nti_check": "Liquid Filling System Specifications - Nozzle type - ibso",
    "lf_ntiq": "Liquid Filling System Specifications - Nozzle type - ibso - qty",
    "lf_nto_check": "Liquid Filling System Specifications - Nozzle type - obso",
    "lf_ntoq": "Liquid Filling System Specifications - Nozzle type - obso - qty",
    "lf_nts_check": "Liquid Filling System Specifications - Nozzle type - straight",
    "lf_ntsq": "Liquid Filling System Specifications - Nozzle type - straight - qty",
    "lf_ntsv_check": "Liquid Filling System Specifications - Nozzle type - straight with check valve",
    "lf_ntsvq": "Liquid Filling System Specifications - Nozzle type - straight with check valve - qty",
    "lf_ocw_check": "Liquid Filling System Specifications - Options - check weighing (tare in/out)",
    "lf_onc": "Liquid Filling System Specifications - Options - no. of cells",
    "lf_p1000": "Liquid Filling System Specifications - Pump - 1000cc (qty)",
    "lf_p100": "Liquid Filling System Specifications - Pump - 100cc (qty)",
    "lf_p10": "Liquid Filling System Specifications - Pump - 10cc (qty)",
    "lf_p250": "Liquid Filling System Specifications - Pump - 250cc (qty)",
    "lf_p500": "Liquid Filling System Specifications - Pump - 500cc (qty)",
    "lf_p50": "Liquid Filling System Specifications - Pump - 50cc (qty)",
    "lf_p850": "Liquid Filling System Specifications - Pump - 850cc (qty)",
    "lf_phc_check": "Liquid Filling System Specifications - Pump - hard chrome",
    "lf_pmm": "Liquid Filling System Specifications - Pump - mass meter (qty)",
    "lf_pnone_check": "Liquid Filling System Specifications - Pump - none",
    "lf_pother": "Liquid Filling System Specifications - Pump - Other",
    "lf_ptl_check": "Liquid Filling System Specifications - Pump type - liquid/semi",
    "lf_ptp_check": "Liquid Filling System Specifications - Pump type - pneumatic",
    "lf_pts_check": "Liquid Filling System Specifications - Pump type - servo (qty)",
    "lf_ptsb_check": "Liquid Filling System Specifications - Pump type - servo bottom-up fill",
    "lf_ptsh_check": "Liquid Filling System Specifications - Pump type - steel heart (peristaltic)",
    "lf_ptv_check": "Liquid Filling System Specifications - Pump type - volumetric",
    "lf_ptvi_check": "Liquid Filling System Specifications - Pump type - viscous",
    "lf_ptpd_check": "Liquid Filling System Specifications - Pump type - positive displacement",
    "lf_vta_check": "Liquid Filling System Specifications - Valve type - air pilot",
    "lf_vtbw_check": "Liquid Filling System Specifications - Valve type - ball weight",
    "lf_vtnone_check": "Liquid Filling System Specifications - Valve type - none",
    "lf_vtr_check": "Liquid Filling System Specifications - Valve type - rotary",
    "lf_cdc_check": "Liquid Filling System Specifications - Centering device - cone",
    "lf_cdn_check": "Liquid Filling System Specifications - Centering device - neck",
    "lf_cdnone_check": "Liquid Filling System Specifications - Centering device - none",
    "lf_th100_check": "Liquid Filling System Specifications - Tank - 100l",
    "lf_th18_check": "Liquid Filling System Specifications - Tank - 18l",
    "lf_th60_check": "Liquid Filling System Specifications - Tank - 60l",
    "lf_tha_check": "Liquid Filling System Specifications - Tank - agitator",
    "lf_thchz": "Liquid Filling System Specifications - Tank - customer hose size (id) or tri-clamp",
    "lf_the_check": "Liquid Filling System Specifications - Tank - electropolished",
    "lf_thh_check": "Liquid Filling System Specifications - Tank - hopper",
    "lf_thitp_check": "Liquid Filling System Specifications - Tank - incl. transfer pump",
    "lf_thj_check": "Liquid Filling System Specifications - Tank - jacketed",
    "lf_thnone_check": "Liquid Filling System Specifications - Tank - none",
    "lf_tho": "Liquid Filling System Specifications - Tank - other",
    "lf_thsb_check": "Liquid Filling System Specifications - Tank - spray ball",
    "lf_tht_check": "Liquid Filling System Specifications - Tank - tank",

    # Street Fighter Tablet Counter mappings
    "sf_100_check": "Street Fighter Tablet Counter - Street fighter - 100",
    "sf_1_check": "Street Fighter Tablet Counter - Street fighter - 1",
    "sf_2_check": "Street Fighter Tablet Counter - Street fighter - 2",
    "sf_lf100_check": "Street Fighter Tablet Counter - Lift fighter - 100l",
    "sf_lf_check": "Street Fighter Tablet Counter - Lift fighter - sf_lf_check",
    "sf_lfi_check": "Street Fighter Tablet Counter - Lift fighter - interlocked",
    "sf_lflc_check": "Street Fighter Tablet Counter - Lift fighter - load cells",
    "sf_lfna_check": "Street Fighter Tablet Counter - Lift fighter - no air",
    "sf_lftah_check": "Street Fighter Tablet Counter - Lift fighter - twin axis hmi",
    "sf_nf1_check": "Street Fighter Tablet Counter - No. of funnels - 1",
    "sf_nf2_check": "Street Fighter Tablet Counter - No. of funnels - 2",
    "sf_nf5_check": "Street Fighter Tablet Counter - No. of funnels - 5",
    "sf_none_check": "Street Fighter Tablet Counter - None",
    "sf_cs_check": "Street Fighter Tablet Counter - Cleaning stn",
    "sf_de_check": "Street Fighter Tablet Counter - Hopper - dust ext.",
    "sf_h10_check": "Street Fighter Tablet Counter - Hopper - 10l",
    "sf_h60_check": "Street Fighter Tablet Counter - Hopper - 60l",

    # Reject / Inspection System mappings
    "rj_bc_check": "Reject / Inspection System - Inspect for - bar code",
    "rj_ch_check": "Reject / Inspection System - Reject method - chute",
    "rj_comm": "Reject / Inspection System - comments",
    "rj_cp_check": "Reject / Inspection System - Inspect for - cap prs.",
    "rj_hc_check": "Reject / Inspection System - Inspect for - skewed cap",
    "rj_hs_check": "Reject / Inspection System - Reject method - high speed",
    "rj_lk_check": "Reject / Inspection System - Reject method - lockable",
    "rj_lp_check": "Reject / Inspection System - Inspect for - label prs",
    "rj_lpo_check": "Reject / Inspection System - Inspect for - label position",
    "rj_lr_check": "Reject / Inspection System - Reject method - linear",
    "rj_nf_check": "Reject / Inspection System - Inspect for - no foil",
    "rj_op_check": "Reject / Inspection System - Inspect for - overcap prs",
    "rj_or_check": "Reject / Inspection System - Inspect for - ocr",
    "rj_ov_check": "Reject / Inspection System - Inspect for - ocv",
    "rj_pp_check": "Reject / Inspection System - Inspect for - pintle prs",
    "rj_sl_check": "Reject / Inspection System - Reject method - starwheel",
    "rj_sp_check": "Reject / Inspection System - Inspect for - stem prs",
    "rj_tk_check": "Reject / Inspection System - Reject method - track",
    "rj_tq_check": "Reject / Inspection System - Inspect for - torque",
    "rj_ty_check": "Reject / Inspection System - Reject method - tray",
    "rj_con_check": "Conveyor Specifications - Reject method - conveyor",

    # Turn Tables mappings
    "tb_39_check": "Turn Tables - 39\"",
    "tb_48_check": "Turn Tables - 48\"",
    "tb_60_check": "Turn Tables - 60\"",
    "tb_buf_check": "Turn Tables - buffer",
    "tb_iwot_check": "Turn Tables - infeed w/o tray",
    "tb_iwt_check": "Turn Tables - infeed table w/tray",
    "tb_or_check": "Turn Tables - orientor",
    "tb_owot_check": "Turn Tables - outfeed w/o tray",
    "tb_owt_check": "Turn Tables - outfeed table w/tray",

    # Control & Programming Specifications mappings
    "batch_at_check": "Control & Programming Specifications - Batch / Data Report - Audit Trail",
    "batch_none_check": "Control & Programming Specifications - Batch / Data Report - None",
    "batch_sht_check": "Control & Programming Specifications - Batch / Data Report - Summary Header with Tracking",
    "batch_yes15_check": "Control & Programming Specifications - Batch / Data Report - Yes (requires 15\" HMI)",
    "blt_audible_check": "Control & Programming Specifications - Beacon Light Tower - Audible",
    "blt_green_check": "Control & Programming Specifications - Beacon Light Tower - Green",
    "blt_none_check": "Control & Programming Specifications - Beacon Light Tower - None",
    "blt_red_check": "Control & Programming Specifications - Beacon Light Tower - Red",
    "blt_yellow_check": "Control & Programming Specifications - Beacon Light Tower - Yellow",
    "cpp_1axis_check": "Control & Programming Specifications - Control Panel Post - 1 Axis",
    "cpp_2axis_check": "Control & Programming Specifications - Control Panel Post - 2 Axis",
    "cpp_3axis_check": "Control & Programming Specifications - Control Panel Post - 3 Axis",
    "cpp_fixed_check": "Control & Programming Specifications - Control Panel Post - Fixed (STD for Explosive Environment)",
    "cps_ep_check": "Control & Programming Specifications - Explosion proof",
    "cps_none_check": "Control & Programming Specifications - Explosion proof - none",
    "etr_10hmi_check": "Control & Programming Specifications - Electronic torque readout - Size - 10\"",
    "etr_none_check": "Control & Programming Specifications - Electronic torque readout - none",
    "hmi_allenb_check": "Control & Programming Specifications - HMI - Allen Bradley",
    "hmi_b&r_check": "Control & Programming Specifications - HMI - B & R",
    "hmi_pc_check": "Control & Programming Specifications - HMI - PC Upgrade",
    "hmi_pv10_check": "Control & Programming Specifications - HMI - Allen Bradley - Size - 10\"",
    "hmi_pv7_check": "Control & Programming Specifications - HMI - Allen Bradley - Size - 7\"",
    "hmi_size10_check": "Control & Programming Specifications - HMI - B & R - Size - 10\"",
    "hmi_size15_check": "Control & Programming Specifications - HMI - PC upgrade - Size - 15\"",
    "hmi_size5.7_check": "Control & Programming Specifications - HMI - B & R - Size - 5.7\" n/a for vision",
    "lan_e_check": "Control & Programming Specifications - HMI - Language - English",
    "lan_f_check": "Control & Programming Specifications - HMI - Language - French",
    "plc_allenb_check": "Control & Programming Specifications - PLC - Allen Bradley",
    "plc_b&r_check": "Control & Programming Specifications - PLC - B & R",
    "plc_compactl_check": "Control & Programming Specifications - PLC - CompactLogix",
    "plc_controll_check": "Control & Programming Specifications - PLC - ControlLogix",
    "rts_co_check": "Control & Programming Specifications - connection only",
    "rts_none_check": "Control & Programming Specifications - none",
    "rts_secomea_check": "Control & Programming Specifications - remote technical service (secomea)",

    # Change Part Quantities and Construction Materials mappings
    "plug_ss304_check": "Plugging System Specifications - ss 304",
    "plug_ss316_check": "Plugging System Specifications - ss 316",

    # Order Identification and Basic Information mappings
    "machine": "Basic Information - Machine",
    "production_speed": "Order Identification - Production speed",
    "quote": "Order Identification - Quote",

    # Utility Specifications mappings
    "amps": "Utility Specifications - AMPS",
    "ce_c1d2_check": "Utility Specifications - Certification - Class1 Div 2",
    "ce_csa_check": "Utility Specifications - Certification - CSA",
    "ce_expl_check": "Utility Specifications - Certification - Explosion",
    "cfm": "Utility Specifications - CFM",
    "conformity_csa_check": "Utility Specifications - csa",
    "country": "Utility Specifications - country",
    "hz": "Utility Specifications - Hz",
    "phases": "Utility Specifications - phases",
    "psi": "Utility Specifications - PSI",
    "voltage": "Utility Specifications - Voltage",

    "options_listing": "Option Listing - Additional Quoted Options"
}

# Define explicit placeholder mappings for SortStar
SORTSTAR_EXPLICIT_MAPPINGS = {
    # Basic Information mappings
    "customer": "GENERAL ORDER ACKNOWLEDGEMENT > Customer",
    "machine": "GENERAL ORDER ACKNOWLEDGEMENT > Machine",
    "direction": "GENERAL ORDER ACKNOWLEDGEMENT > Direction",
    "quote": "Order Identification > Quote",
    "production_speed": "Order Identification > Production speed",
    #Utility Specifications mappings
    "voltage": "Utility Specifications > Voltage",
    "phases": "Utility Specifications > Phases",
    "hz": "Utility Specifications > Hz",
    "amps": "Utility Specifications > AMPS",
    "psi": "Utility Specifications > PSI",
    "cfm": "Utility Specifications > CFM",
    "conformity_csa_check": "Utility Specifications > Conformity CSA Check",
    "ce_none_check": "Utility Specifications > CE None Check",
    "ce_csa_check": "Utility Specifications > CE CSA Check",
    "ce_expl_check": "Utility Specifications > CE Expl Check",
    "country": "Utility Specifications > Country of destination",
    # Basic Systems mappings
    "bs_984_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 220VAC 3 Phases LEFT TO RIGHT",
    "bs_1230_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 220VAC 3 Phases RIGHT TO LEFT",
    "bs_985_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 480VAC & 380VAC 3 Phases LEFT TO RIGHT",
    "bs_1229_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 480VAC & 380VAC 3 Phases RIGHT TO LEFT",
    "bs_1264_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 24ft3 220VAC 3 Phases LEFT TO RIGHT",
    "bs_1265_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 24ft3 480VAC & 380VAC 3 Phases LEFT TO RIGHT",
    # Optional Systems mappings
    "op_2409_check": "OPTIONAL SYSTEMS > Guarding System > Lexan Euroguard Top Cover",
    "op_ias_check": "OPTIONAL SYSTEMS > Guarding System > Ionized Air System",
    "op_nrck_check": "OPTIONAL SYSTEMS > Guarding System > Requires Cap Kit",
    "el_0369_check": "OPTIONAL SYSTEMS > Electrical > Stack Light with Buzzer Option",
    "cps_ep_check": "OPTIONAL SYSTEMS > Control Specifications > Explosion proof",
    "cps_none_check": "OPTIONAL SYSTEMS > Control Specifications > None",
    "plc_b&r_check": "OPTIONAL SYSTEMS > Control Specifications > PLC B & R",
    "plc_compactl_check": "OPTIONAL SYSTEMS > Control Specifications > CompactLogix",
    "plc_controll_check": "OPTIONAL SYSTEMS > Control Specifications > ControlLogix",
    "hmi_b&r_check": "OPTIONAL SYSTEMS > Control Specifications > HMI B & R",
    "hmi_allenb_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Allen Bradley",
    "hmi_pc_check": "OPTIONAL SYSTEMS > Control Specifications > HMI PC Upgrade",
    "hmi_size5.7_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Size 5.7\"",
    "hmi_size10_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Size 10\"",
    "cpp_1axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 1 Axis",
    "cpp_2axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 2 Axis – U-shaped",
    "cpp_3axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 3 Axis",
    "rts_secomea_check": "OPTIONAL SYSTEMS > Remote Technical Service (Secomea)",
    "rts_none_check": "OPTIONAL SYSTEMS > Remote Technical Service > None",
    "rts_co_check": "OPTIONAL SYSTEMS > Remote Technical Service > Connection Only",
    "eg_none_check": "OPTIONAL SYSTEMS > Euro guarding > None",
    "eg_pnl_check": "OPTIONAL SYSTEMS > Euro guarding > Panel material Lexan",
    "eg_pmtg_check": "OPTIONAL SYSTEMS > Euro guarding > Tempered glass",
    "eg_stkw_check": "OPTIONAL SYSTEMS > Euro guarding > Switch type Key switch",
    "eg_stm_check": "OPTIONAL SYSTEMS > Euro guarding > Switch type Magnetic",
    "eg_sto_check": "OPTIONAL SYSTEMS > Euro guarding > Other",
    "eg_rc_check": "OPTIONAL SYSTEMS > Euro guarding > Reject Cover",
    "eg_rcnone_check": "OPTIONAL SYSTEMS > Euro guarding > Reject Cover None",
    "vd_f_check": "OPTIONAL SYSTEMS > Validation Documents > FAT",
    "vd_s_check": "OPTIONAL SYSTEMS > Validation Documents > SAT",
    "vd_d_check": "OPTIONAL SYSTEMS > Validation Documents > DQ",
    "vd_h_check": "OPTIONAL SYSTEMS > Validation Documents > HDS/SDS",
    "vd_fd_check": "OPTIONAL SYSTEMS > Validation Documents > FS/DS",
    "vd_i_check": "OPTIONAL SYSTEMS > Validation Documents > IQ/OQ",
    "pt_tcc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Transport charges Capmatic",
    "pt_tccu_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Transport charges Customer",
    "pt_pbc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging by Capmatic",
    "pt_pbcu_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging by Customer / Not Incl.",
    "pt_pts_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Skid",
    "pt_ptc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Crate",
    "pt_ptsw_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Sea Worthy",
    "wi_war1_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Warranty 1YR",
    "wi_war2_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Warranty 2YR",
    "sk_none_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit None",
    "sk_1yr_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit 1YR",
    "sk_2yr_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit 2YR",
    "wrts_y_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Remote Tech. Service Yes",
    "stpc_none_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Start-up Commissioning None",
    "stpc_yes": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Start-up Commissioning Yes – no. of days",
    "options_listing": "Option Listing > Additional Quoted Options"
}

def extract_placeholders(template_path: str) -> List[str]:
    """Reads the template and extracts all unique, cleaned placeholder keys."""
    print(f"Extracting placeholders from: {template_path}")
    placeholders: Set[str] = set()
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")

        for para in doc.paragraphs:
            for match in regex.findall(para.text):
                cleaned_key = match.strip()
                if cleaned_key:
                    placeholders.add(cleaned_key)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in regex.findall(cell.text):
                        cleaned_key = match.strip()
                        if cleaned_key:
                            placeholders.add(cleaned_key)
                            
        if not placeholders:
             print(f"Warning: No placeholders found in {template_path}")
             return []

        print(f"Found {len(placeholders)} unique placeholders.")
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error reading placeholders from template '{template_path}': {e}")
        return []

def is_likely_section_header(paragraph) -> bool:
    """Heuristically determines if a paragraph is a section header."""
    text = paragraph.text.strip()
    if not text: # Empty paragraphs are not headers
        return False
    
    # Condition 1: All caps and relatively short
    if text.isupper() and len(text.split()) < 7:
        return True
        
    # Condition 2: Ends with a colon and is relatively short
    # if text.endswith(":") and len(text.split()) < 7:
    #     return True # This might be too broad, catching field labels

    # Condition 3: Bold text (check first run)
    if paragraph.runs and paragraph.runs[0].bold and len(text.split()) < 7:
        # Further check: not too many lowercase letters if it's mostly uppercase (e.g. avoids bolded sentences)
        if sum(1 for char in text if char.islower()) < len(text) / 2:
             return True
    
    # Condition 4: Check for specific heading styles if used consistently in template
    # for style_name_part in ["heading 1", "heading 2", "heading 3", "title"]:
    #     if paragraph.style and paragraph.style.name and style_name_part in paragraph.style.name.lower():
    #         return True
            
    # Add more heuristics as needed (e.g., font size significantly larger than body text)
    return False

def extract_placeholder_context_hierarchical(template_path: str, 
                                            explicit_placeholder_mappings: Dict[str, str],
                                            enhance_with_outline: bool = True,
                                            outline_path: str = "full_fields_outline.md",
                                            check_if_all_mapped: bool = True,  # Changed back to True as default
                                            is_sortstar: bool = False) -> Dict[str, str]:
    """
    Parses the template to extract placeholders and attempts to build hierarchical context
    by identifying section headers. Optionally enhances with outline file.
    
    Args:
        template_path: Path to the Word document template
        explicit_placeholder_mappings: The explicit mappings for the given template type.
        enhance_with_outline: Whether to enhance context with outline file
        outline_path: Path to the outline file
        check_if_all_mapped: If True, uses explicit mappings as base and enhances with dynamic extraction
        is_sortstar: Flag to handle SortStar specific logic.
    """
    print(f"Extracting hierarchical placeholder context from: {template_path}")
    
    # Start with explicit mappings as the base
    context_map: Dict[str, str] = {}
    if explicit_placeholder_mappings:
        context_map.update(explicit_placeholder_mappings)
        print(f"Using {len(explicit_placeholder_mappings)} explicit mappings as base")
    
    # Get all placeholders from the template
    all_placeholders = extract_placeholders(template_path)
    if not all_placeholders:
        print("No placeholders found in template")
        return context_map
        
    # If check_if_all_mapped is True, we'll still do dynamic extraction but use it to enhance existing mappings
    if check_if_all_mapped:
        unmapped = [ph for ph in all_placeholders if ph not in explicit_placeholder_mappings]
        if unmapped:
            print(f"Found {len(unmapped)} unmapped placeholders. Will enhance with dynamic extraction.")
            if len(unmapped) <= 10:
                print(f"Unmapped placeholders: {', '.join(unmapped)}")
        else:
            print("All placeholders are explicitly mapped. Will still perform dynamic extraction for enhancement.")
    
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")
        
        current_section_header = "General" # Default section
        current_subsection_header = ""

        # First pass: Identify all placeholders and their immediate cell/paragraph text
        placeholder_details = {}

        for p_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if is_likely_section_header(para):
                current_section_header = para_text.replace(":","").strip()
                current_subsection_header = ""
                continue
            
            elif para.runs and para.runs[0].bold and len(para_text.split()) < 5 and not para_text.isupper() and para_text.endswith(":"):
                current_subsection_header = para_text.replace(":","").strip()
                continue

            for r_match in regex.finditer(para_text):
                ph_key = r_match.group(1).strip()
                
                # Always process the placeholder, but respect existing mappings
                if ph_key and ph_key not in placeholder_details:
                    preceding_text = para_text[:r_match.start()].strip()
                    preceding_text = regex.sub("", preceding_text).strip().replace(":","").strip()
                    placeholder_details[ph_key] = {
                        "immediate_label": preceding_text if preceding_text else ph_key,
                        "section": current_section_header,
                        "subsection": current_subsection_header,
                        "type": "paragraph"
                    }

        # Process tables similarly...
        for t_idx, table in enumerate(doc.tables):
            current_table_group_label = ""
            for r_idx, row in enumerate(table.rows):
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text and not regex.search(first_cell_text) and len(first_cell_text.split()) < 4:
                        if r_idx == 0 or (r_idx > 0 and table.cell(r_idx,0).text != table.cell(r_idx-1,0).text):
                            current_table_group_label = first_cell_text.replace(":","").strip()
                        elif not current_table_group_label:
                            current_table_group_label = first_cell_text.replace(":","").strip()
                
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for r_match in regex.finditer(cell_text):
                        ph_key = r_match.group(1).strip()

                        if ph_key and ph_key not in placeholder_details:
                            immediate_label = ""
                            if c_idx > 0:
                                label_cell_text = row.cells[c_idx-1].text.strip().replace(":","").strip()
                                if label_cell_text and not regex.search(label_cell_text):
                                    immediate_label = label_cell_text
                            if not immediate_label:
                                immediate_label = cell_text[:r_match.start()].strip().replace(":","").strip()
                            
                            placeholder_details[ph_key] = {
                                "immediate_label": immediate_label if immediate_label else ph_key,
                                "section": current_section_header,
                                "subsection": current_subsection_header,
                                "table_group": current_table_group_label if current_table_group_label != immediate_label else "",
                                "type": "table"
                            }

        # Enhance existing mappings with dynamic context
        for ph_key, details in placeholder_details.items():
            parts = []
            if details.get("section") and details["section"] != "General": parts.append(details["section"])
            if details.get("subsection"): parts.append(details["subsection"])
            if details.get("table_group"): parts.append(details["table_group"])
            if details.get("immediate_label") and details["immediate_label"] != ph_key: parts.append(details["immediate_label"])
            
            if parts:
                dynamic_context = " - ".join(filter(None, parts))
                # If we have an existing mapping, enhance it with dynamic context if it adds value
                if ph_key in context_map:
                    existing_context = context_map[ph_key]
                    # Only update if dynamic context provides additional information
                    if len(dynamic_context) > len(existing_context) or " - " in dynamic_context:
                        context_map[ph_key] = dynamic_context
                else:
                    context_map[ph_key] = dynamic_context

        # Ensure all placeholders have some context
        for ph in all_placeholders:
            if ph not in context_map:
                context_map[ph] = ph

        if not context_map: 
            print(f"Warning: No placeholder context generated for {template_path}")
        else: 
            print(f"Generated hierarchical context for {len(context_map)} placeholders.")
        
        # Enhance with outline file if requested
        if enhance_with_outline and os.path.exists(outline_path) and not is_sortstar:
            context_map = enhance_placeholder_context_with_outline(context_map, explicit_placeholder_mappings, outline_path)
        
        return context_map

    except Exception as e:
        print(f"Error extracting hierarchical placeholder context from '{template_path}': {e}")
        import traceback
        traceback.print_exc()
        return {}

def extract_placeholder_schema(template_path: str, explicit_mappings: Optional[Dict[str, str]] = None, is_sortstar: bool = False) -> Dict[str, Dict]:
    """
    Creates a structured JSON schema from the template with rich metadata about each field.
    
    Args:
        template_path: The path to the Word document template.
        explicit_mappings: The explicit mappings for the given template type.
        is_sortstar: Flag to handle SortStar specific logic.

    Returns a dictionary where:
    - Each key is a placeholder name
    - Each value is a dictionary with:
        - type: "string" or "boolean" (for _check fields)
        - section: The main section this field belongs to
        - subsection: The subsection if applicable
        - description: A human-readable description of the field
        - location: Where in the document the field appears (paragraph or table)
        - synonyms: List of alternative terms (for checkboxes)
        - positive_indicators: List of phrases that indicate the checkbox should be YES
        
    This schema is designed to be used by LLMs to better understand the template structure.
    """
    print(f"Extracting JSON schema from template: {template_path}")
    
    # First, use the existing function to get the placeholder details
    placeholder_details = {}
    schema = {}

    # Populate schema with explicitly mapped placeholders first if they are provided
    if explicit_mappings:
        for ph_key, description in explicit_mappings.items():
            field_type = "boolean" if ph_key.endswith("_check") else "string"
            # Basic section/subsection parsing from description string
            delimiter = " > " if is_sortstar else " - "
            parts = description.split(delimiter)
            section = parts[0] if len(parts) > 0 else "General"
            subsection = parts[1] if len(parts) > 1 else ""
            
            schema[ph_key] = {
                "type": field_type,
                "section": section,
                "subsection": subsection,
                "description": description, # Use full path as description
                "location": "unknown" # Location is harder to determine from flat list
            }
            if field_type == "boolean":
                # Simplified synonym/indicator generation for explicitly mapped items
                synonyms = [description.lower(), ph_key.replace("_check","").replace("_"," ")]
                schema[ph_key]["synonyms"] = list(set(synonyms))
                schema[ph_key]["positive_indicators"] = [f"with {s}" for s in synonyms] + ["yes", "selected", description.lower()]
    
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")
        
        current_section_header = "General" # Default section
        current_subsection_header = ""

        # First pass: Identify all placeholders and their immediate cell/paragraph text
        for p_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if is_likely_section_header(para):
                current_section_header = para_text.replace(":","").strip()
                current_subsection_header = "" # Reset subsection on new main section
                continue
            
            # Heuristic for subsection: If text is bold and short but not ALL CAPS
            elif para.runs and para.runs[0].bold and len(para_text.split()) < 5 and not para_text.isupper() and para_text.endswith(":"):
                current_subsection_header = para_text.replace(":","").strip()
                continue

            for r_match in regex.finditer(para_text):
                ph_key = r_match.group(1).strip()
                if ph_key and ph_key not in schema: # Only process if not already in schema from explicit map
                    preceding_text = para_text[:r_match.start()].strip()
                    preceding_text = regex.sub("", preceding_text).strip().replace(":","").strip()
                    placeholder_details[ph_key] = {
                        "immediate_label": preceding_text if preceding_text else ph_key,
                        "section": current_section_header,
                        "subsection": current_subsection_header,
                        "type": "paragraph"
                    }

        for t_idx, table in enumerate(doc.tables):
            current_table_group_label = "" # For labels like HMI, PLC in the first column spanning rows
            for r_idx, row in enumerate(table.rows):
                # Try to detect a group label in the first column if it spans or is consistent
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text and not regex.search(first_cell_text) and len(first_cell_text.split()) < 4:
                        # Heuristic: if this cell is different from the one above it in the same column, 
                        # or if it's the first row, it might be a new group label for subsequent rows.
                        if r_idx == 0 or (r_idx > 0 and table.cell(r_idx,0).text != table.cell(r_idx-1,0).text):
                            current_table_group_label = first_cell_text.replace(":","").strip()
                        elif not current_table_group_label: # If still no group label, take from first row, first cell
                            current_table_group_label = first_cell_text.replace(":","").strip()
                
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for r_match in regex.finditer(cell_text):
                        ph_key = r_match.group(1).strip()
                        if ph_key and ph_key not in schema: # Only process if not already in schema from explicit map
                            immediate_label = ""
                            if c_idx > 0: # Label to the left
                                label_cell_text = row.cells[c_idx-1].text.strip().replace(":","").strip()
                                if label_cell_text and not regex.search(label_cell_text):
                                    immediate_label = label_cell_text
                            if not immediate_label: # Text before in current cell
                                immediate_label = cell_text[:r_match.start()].strip().replace(":","").strip()
                            
                            placeholder_details[ph_key] = {
                                "immediate_label": immediate_label if immediate_label else ph_key,
                                "section": current_section_header,
                                "subsection": current_subsection_header, # Could be overridden by table group
                                "table_group": current_table_group_label if current_table_group_label != immediate_label else "",
                                "type": "table"
                            }

        # Convert the details to a structured schema
        for ph_key, details in placeholder_details.items():
            field_type = "boolean" if ph_key.endswith("_check") else "string"
            
            # Build a descriptive label combining all context
            description_parts = []
            if details.get("immediate_label") and details["immediate_label"] != ph_key:
                description_parts.append(details["immediate_label"])
            if details.get("table_group"):
                description_parts.append(details["table_group"])
                
            description = " - ".join(filter(None, description_parts)) or ph_key
                
            # Create the schema entry
            schema[ph_key] = {
                "type": field_type,
                "section": details.get("section", "General"),
                "subsection": details.get("subsection", ""),
                "description": description,
                "location": details.get("type", "unknown")
            }
            
            # For checkboxes, add helpful metadata with synonyms and positive indicators
            if field_type == "boolean":
                # Generate common synonyms for this field based on the key and description
                synonyms = generate_synonyms_for_checkbox(ph_key, description)
                schema[ph_key]["synonyms"] = synonyms
                
                # Generate positive indicators based on the synonyms and key
                positive_indicators = generate_positive_indicators(ph_key, description, synonyms)
                schema[ph_key]["positive_indicators"] = positive_indicators
            
        # Ensure all placeholders have schema entries
        all_phs = extract_placeholders(template_path)
        for ph in all_phs:
            if ph not in schema:
                field_type = "boolean" if ph.endswith("_check") else "string"
                schema[ph] = {
                    "type": field_type,
                    "section": "General",
                    "subsection": "",
                    "description": ph,
                    "location": "unknown"
                }
                
                # For checkboxes, add helpful metadata
                if field_type == "boolean":
                    synonyms = generate_synonyms_for_checkbox(ph, ph)
                    schema[ph]["synonyms"] = synonyms
                    positive_indicators = generate_positive_indicators(ph, ph, synonyms)
                    schema[ph]["positive_indicators"] = positive_indicators

        print(f"Generated schema for {len(schema)} placeholders.")
        return schema

    except Exception as e:
        print(f"Error extracting placeholder schema from '{template_path}': {e}")
        import traceback
        traceback.print_exc()
        return {}

def generate_synonyms_for_checkbox(key: str, description: str) -> List[str]:
    """
    Generates synonyms for a checkbox field based on its key and description.
    
    Args:
        key: The placeholder key (e.g. "explosion_proof_check")
        description: The human-readable description
        
    Returns:
        A list of synonyms for this concept
    """
    # Remove _check suffix for cleaner key
    clean_key = key.replace("_check", "")
    
    # Comprehensive dictionary of packaging industry terms and their synonyms
    word_synonyms = {
        # Safety and compliance terms
        "explosion_proof": ["explosion-proof", "explosion proof", "explosion protected", "exp", "exd", "class 1 div 2", 
                           "explosion protection", "hazardous area", "atex", "ex-proof", "explosion protected environment"],
        "stainless": ["ss", "s.s.", "stainless steel", "inox", "s/s", "316", "304", "316l", "ss316", "ss304", "aisi"],
        "certification": ["certified", "ce", "csa", "ul", "ansi", "asme", "iso", "iec", "din", "gmp", "fda", "3a"],
        "gmp": ["good manufacturing practice", "gmp-compliant", "cleanroom", "clean room", "pharmaceutical grade"],
        "fda": ["food grade", "fda approved", "fda-compliant", "pharma grade", "medical grade"],
        
        # Production parameters
        "production_speed": ["projected speed", "throughput", "bottles per minute", "units per minute", "parts per minute", 
                            "bpm", "upm", "ppm", "production rate", "processing speed", "output rate", "machine speed"],
        
        # Identification and coding systems
        "barcode": ["bar code", "barcode reader", "code reader", "scanner", "scan", "scanning", "2d code", "qr", 
                   "datamatrix", "data matrix", "upc", "ean", "barcode verification", "symbology"],
        "ocr": ["optical character recognition", "character reading", "text reading", "code reading", "character verification"],
        "ocv": ["optical character verification", "text verification", "character validation", "text validation"],
        "vision": ["machine vision", "camera", "imaging", "visual inspection", "vision system", "cognex", "keyence", "omron"],
        "laser": ["laser marking", "laser coding", "laser etching", "laser printing", "co2 laser", "fiber laser", "marking laser"],
        "inkjet": ["ink jet", "ink-jet", "cij", "continuous inkjet", "thermal inkjet", "ink printer", "coding", "printing"],
        
        # Labeling and packaging systems
        "label": ["label", "labelling", "labeling", "labeler", "labeller", "adhesive label", "pressure sensitive", 
                 "roll fed", "cut and stack", "label applicator", "front label", "back label"],
        "sleeve": ["shrink sleeve", "shrink label", "sleeve applicator", "full body sleeve", "neck sleeve", "tamper evident"],
        "wrap": ["wrap around", "wraparound", "wrap-around", "wrap label", "full wrap", "partial wrap"],
        "reel": ["label reel", "roll", "spool", "material reel", "unwinder", "rewinder", "roll holder"],
        
        # Control systems
        "hmi": ["hmi", "human machine interface", "touch screen", "touch panel", "operator interface", 
               "control panel", "display", "panel pc", "operator panel", "touchscreen", "monitor", "5.7 hmi", "5.7 inch hmi"],
        "plc": ["plc", "controller", "control system", "automation controller", "programmable logic", 
               "automation system", "control unit", "processor", "allen bradley", "siemens", "b&r", "omron", "b & r plc", "b&r plc"],
        "servo": ["servo motor", "servo drive", "servo system", "servo control", "brushless", "motion control", 
                 "precision motion", "servo-driven", "servo actuator", "stepper", "motor"],
        "pneumatic": ["air", "pneumatic cylinder", "pneumatic actuator", "pneumatic system", "compressed air", 
                     "air cylinder", "air pressure", "air operated", "air driven"],
        
        # Transport and handling systems
        "conveyor": ["conveying system", "transport system", "belt conveyor", "chain conveyor", "mat conveyor", 
                    "conveyor belt", "transport", "product transfer", "product handling"],
        "accumulation": ["accumulation table", "accumulator", "buffer", "buffer table", "accumulation conveyor", 
                        "bottle accumulator", "container buffer"],
        "elevator": ["product elevator", "vertical conveyor", "cap elevator", "bottle elevator", "vertical transport", 
                    "lifting system", "bucket elevator", "z-elevator"],
        "turntable": ["rotary table", "turn table", "rotating table", "accumulation table", "indexing table", 
                     "disc table", "disc turntable", "rotary buffer"],
        "starwheel": ["star wheel", "timing screw", "timing star", "infeed star", "discharge star", "transfer star", 
                     "pocket wheel", "container transfer"],
        
        # Filling and dispensing systems
        "filling": ["filler", "filling system", "liquid filling", "volumetric filling", "gravimetric filling", 
                   "level filling", "time pressure filling", "mass flow", "piston filler"],
        "pump": ["peristaltic pump", "gear pump", "lobe pump", "piston pump", "diaphragm pump", "centrifugal pump", 
                "rotary pump", "dosing pump", "metering pump", "dispensing pump"],
        "nozzle": ["filling nozzle", "dispensing nozzle", "fill head", "dosing nozzle", "spray nozzle", 
                  "injection nozzle", "applicator nozzle", "valve nozzle"],
        "valve": ["filling valve", "control valve", "check valve", "solenoid valve", "ball valve", "needle valve", 
                 "butterfly valve", "diaphragm valve", "pinch valve"],
        
        # Capping and sealing systems
        "capping": ["capper", "cap applicator", "cap tightener", "cap sealer", "screwing", "twist-off", 
                   "screw capper", "press-on capper", "snap-on capper", "ROPP", "roll-on pilfer-proof"],
        "torque": ["torque control", "cap torque", "torque monitoring", "torque verification", "torque check", 
                  "torque testing", "torque adjustment", "tightening torque"],
        "sealing": ["heat sealing", "induction sealing", "foil sealing", "ultrasonic sealing", "band sealing", 
                   "conduction sealing", "hermetic seal", "tamper evident"],
        "induction": ["induction sealer", "cap sealer", "foil sealer", "induction heating", "induction coil", 
                     "sealing head", "cap sealing"],
        
        # Detection and verification systems
        "sensor": ["detector", "sensing device", "photoelectric", "proximity", "ultrasonic", "capacitive", 
                  "inductive", "fiber optic", "vision sensor", "level sensor", "presence sensor"],
        "inspection": ["inspection system", "quality control", "verification", "checking", "monitoring", 
                      "detection", "visual inspection", "automated inspection", "100% inspection"],
        "reject": ["rejection system", "reject mechanism", "reject station", "rejection station", "ejector", 
                  "rejection device", "sort", "discard", "rejection arm"],
        "verification": ["check", "verification system", "monitoring", "quality assurance", "validation", 
                        "confirmation", "inspection", "testing", "quality check"],
        
        # Container handling terms
        "bottle": ["container", "vial", "jar", "flask", "ampoule", "bottle", "can", "packaging", "primary container"],
        "puck": ["carrier", "container carrier", "bottle puck", "vial carrier", "nest", "container holder", "pocket"],
        "unscrambler": ["bottle unscrambler", "container unscrambler", "container orienter", "bottle orienter", 
                       "unscrambling system", "bottle sorting", "bottle feed"],
        "accumulator": ["buffer", "accumulation", "accumulation table", "container buffer", "bottle buffer", 
                       "buffering system", "product queue"],
        
        # Machine features
        "clean_in_place": ["cip", "clean-in-place", "cip system", "automated cleaning", "washdown", "cleaning system",
                          "sanitization", "sterilization", "aseptic", "clean room"],
        "touchless": ["no-touch", "non-contact", "contactless", "touchless operation", "hands-free", "remote operation"],
        "remote_service": ["remote support", "remote access", "remote diagnostics", "remote monitoring", 
                          "teleservice", "remote maintenance", "remote connection"],
        "guarding": ["safety guarding", "machine guarding", "safety fence", "protective cover", "safety barrier", 
                    "safety shield", "protective enclosure", "lexan", "plexiglass", "safety door"],
        
        # Specialized add-ons
        "weighing": ["checkweigher", "weight verification", "weight control", "gravimetric", "scale", "balance", 
                    "load cell", "mass measurement", "weight check"],
        "coding": ["printer", "marking", "coding system", "date coder", "lot coder", "batch coder", 
                  "expiry date printer", "variable information"],
        "serialization": ["track and trace", "unique identification", "serial number", "unique code", 
                         "serialization system", "aggregation", "parent-child relationship"],
        "format_change": ["changeover", "format parts", "change parts", "size parts", "adjustment parts", 
                         "quick change", "tool-less changeover", "format conversion"]
    }
    
    # Extract potential key parts by splitting on underscores
    key_parts = clean_key.split("_")
    
    # Get synonyms from description
    desc_words = description.lower().split()
    
    # Start with the key itself and the description
    synonyms = [key, clean_key.replace("_", " ")]

    # Add variations from description
    if description:
        desc_lower = description.lower()
        synonyms.append(desc_lower) # Raw lowercase description

        # Description with punctuation (except period and ampersand) removed, spaces kept
        desc_no_punct_spaces = re.sub(r'[^\w\s.&]', '', desc_lower) # Keep periods and ampersands
        desc_no_punct_spaces = re.sub(r'\s+', ' ', desc_no_punct_spaces).strip() # Consolidate spaces
        synonyms.append(desc_no_punct_spaces)

        # Description with all non-alphanumeric chars removed (concatenated)
        desc_concatenated = re.sub(r'[\W_]', '', desc_lower)
        synonyms.append(desc_concatenated)
    
    # Add synonyms for each key part if available
    for part in key_parts:
        if part in word_synonyms:
            synonyms.extend(word_synonyms[part])
    
    # Look for known phrases in the description and add their synonyms
    for phrase, syn_list in word_synonyms.items():
        if phrase in description.lower() or phrase.replace("_", " ") in description.lower():
            synonyms.extend(syn_list)
    
    # Clean up and return unique values
    cleaned_synonyms = [s.strip().lower() for s in synonyms if s.strip()]
    return list(set(cleaned_synonyms))

def generate_positive_indicators(key: str, description: str, synonyms: List[str]) -> List[str]:
    """
    Generates phrases that would indicate this checkbox should be marked YES.
    
    Args:
        key: The placeholder key
        description: The human-readable description
        synonyms: List of synonyms for this concept
        
    Returns:
        A list of phrases that would indicate this is selected
    """
    # Start with basic indicators
    indicators = ["included", "standard", "included as standard", "yes", "selected"]
    
    # Create phrases like "with <synonym>"
    for synonym in synonyms:
        if synonym:
            indicators.append(f"with {synonym}")
            indicators.append(f"includes {synonym}")
            indicators.append(f"{synonym} included")
            indicators.append(f"{synonym} is selected")
    
    # Clean up the description to use as an indicator
    if description:
        clean_desc = description.strip().lower()
        indicators.append(clean_desc)
        indicators.append(f"with {clean_desc}")
        indicators.append(f"includes {clean_desc}")
    
    # Clean up and return unique values
    cleaned_indicators = [i.strip().lower() for i in indicators if i.strip()]
    return list(set(cleaned_indicators))

def add_section_aware_instructions(template_schema: Dict[str, Dict], prompt_parts: List[str]) -> List[str]:
    """
    Enhances the prompt with section-specific instructions to improve the LLM's understanding
    of each template section's purpose and meaning.
    
    Args:
        template_schema: The JSON schema generated by extract_placeholder_schema
        prompt_parts: The current prompt parts list to enhance
        
    Returns:
        Enhanced prompt parts list with section-specific instructions
    """
    if not template_schema:
        return prompt_parts  # No schema to work with
    
    # Group fields by section
    sections = {}
    for key, field_info in template_schema.items():
        section = field_info.get("section", "General")
        if section not in sections:
            sections[section] = []
        sections[section].append((key, field_info))
    
    # Add section-specific instructions
    if sections:
        prompt_parts.append("\n## SECTION-SPECIFIC GUIDANCE:")
        prompt_parts.append("Pay careful attention to these section-specific instructions when filling out the template:")
        
        # Domain-specific section guidance from the GOA filling guide
        section_guidance = {
            "Basic Information": "This section contains the fundamental project identifiers including project number (Ax), customer name, machine type, and flow direction (typically left-to-right or right-to-left).",
            
            "Order Identification": "This section includes customer PO number, quote number, internal order number (Ox), customer ID, and production speed. All information here should come directly from the quote/sales order document.",
            
            "Utility Specifications": "Look for electrical and pneumatic requirements: supply voltage (e.g., 208-240V), frequency (Hz), air pressure (PSI), certification standards (CSA, CE, etc.), and destination country. IMPORTANT: For explosion-proof environments, electrical components are replaced with pneumatic equivalents.",
            
            "Change Part Quantities and Construction Materials": "This section defines quantities and materials for components that may need to be changed for different product formats, including bottles, plugs, caps, and materials for seals and tubing.",
            
            "Material Specifications": "Identifies materials that come in contact with the product. Look for FDA-approved materials, product contact surfaces (usually SS 304 or SS 316L based on product compatibility), and options like 'Autoclavable' or 'Electropolished'.",
            
            "Control & Programming Specifications": "This critical section defines the machine's control system including explosion proof requirements, PLC type (B&R, Allen Bradley, CompactLogix, ControlLogix), HMI specifications (size, language, location), control panel configuration, beacon lights, E-stops, and data reporting capabilities.",
            
            "Bottle Handling System Specifications": "Defines how containers are fed into the machine, including tube handling systems, vial/bottle transport mechanisms, puck systems for unstable containers, turntables, and indexing mechanisms (starwheel, walking beam, etc.).",
            
            "Reject / Inspection System": "Specifies how rejected products are handled and what conditions trigger rejection. Look for reject methods (chute, tray, conveyor) and reject reasons (fill weight, cap presence, label position, etc.).",
            
            "Street Fighter Tablet Counter": "Specific to tablet counting systems. Identifies model version, number of funnels, hopper size, and features like dust extraction or load cells.",
            
            "Liquid Filling System Specifications": "For filling machines, identifies pump type (volumetric, peristaltic), filling mechanism, pump volume, valve type, nozzle configuration, and options like check weighing. Note that size-appropriate nozzles are typically 2-3mm smaller than the neck opening.",
            
            "Gas Purge": "For products requiring oxygen removal, specifies gas type (nitrogen, argon) and application points (before fill, at fill, after fill, tunnel).",
            
            "Desiccant": "For systems that insert desiccant, specifies type (roll/pouch, cannister) and feeding mechanism.",
            
            "Cottoner": "For systems that insert cotton, includes sensing options (presence, high) and cotton bin configuration.",
            
            "Plugging System Specifications": "For machines that insert plugs, defines insertion mechanism, sorting method, and bulk feeding system.",
            
            "Capping System Specifications": "For machines that apply caps, details cap placement method, torque mechanism and range, cap sorting, centering device, and bulk feeding system. NOTE: A servo with magnetic clutch is not standard and should be verified.",
            
            "BeltStar System Specifications": "For belt-driven capping systems, includes cap placement method, torque mechanism, cap sorting, and adjustment options.",
            
            "Labeling System Specifications": "For labeling machines, identifies label head model, reel diameter, arm type, application system (wrap around, multi-panel), and separator wheel configuration.",
            
            "Coding and Inspection System Specifications": "For printing batch codes and other variable information, includes coder type (hot stamp, thermal transfer, laser, inkjet), vision system configuration, and print content (lot, barcode, expiration date).",
            
            "Induction Specifications": "For induction sealing systems, specifies model, voltage, frequency, sealing head type, stand configuration, and cap inspection capabilities.",
            
            "Conveyor Specifications": "Details conveyor width, length, height, shape, chain type, bed type, and transfer guides.",
            
            "Euro Guarding": "Specifies machine guarding requirements including panel material, switch type, and covers.",
            
            "Validation Documents": "Identifies required documentation (FAT, SAT, DQ, IQ/OQ) and languages. Includes Factory Acceptance Test (FAT) and Site Acceptance Test (SAT) protocol packages with documentation, support, and customer review. Also covers Design Qualification (DQ), Hardware/Software Design Specifications (HDS/SDS), Functional/Design Specifications (FS/DS), and Installation/Operational Qualification (IQ/OQ) documentation in specified languages.",
            
            "Warranty & Install & Spares": "Specifies warranty period, spare parts kit requirements, and commissioning services."
        }
        
        # Add instructions for each section
        for section, fields in sorted(sections.items()):
            # Skip very small sections or generic sections
            if len(fields) < 2 or section in ["General"]:
                continue
            
            # Add domain-specific guidance from our dictionary if available
            if section in section_guidance:
                prompt_parts.append(f"\nFor the '{section}' section: {section_guidance[section]}")
            else:
                # Default guidance based on section name
                if "customer" in section.lower() or "client" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Focus on extracting customer/client details like company name, address, contact information, etc. Look for this information at the beginning of quotes or in header sections.")
                
                elif "machine" in section.lower() or "equipment" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Focus on the specific machine being quoted. The machine model and specifications are usually prominently featured in the quote's main description or line items.")
                
                elif "feature" in section.lower() or "option" in section.lower() or "accessory" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: These are optional features or add-ons for the machine. Check each line item description carefully to determine which options are included in the quote.")
                
                elif "safety" in section.lower() or "compliance" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Look for safety features and compliance standards mentioned in the quote. These might be in dedicated sections or embedded within feature descriptions.")
                    
                elif "warranty" in section.lower() or "service" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Focus on warranty terms, service plans, and support offerings. These are often mentioned toward the end of quotes or in special sections.")
                    
                elif "delivery" in section.lower() or "shipping" in section.lower() or "transport" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Extract information about delivery timelines, shipping methods, and transportation details. Check both line items and any terms & conditions sections.")
                    
                elif "payment" in section.lower() or "financial" in section.lower() or "price" in section.lower():
                    prompt_parts.append(f"For the '{section}' section: Look for payment terms, financial arrangements, and pricing details. These may be in dedicated sections or near the end of the quote.")
                
                # If no specific rule matched but there are fields, create a generic one
                else:
                    prompt_parts.append(f"For the '{section}' section: Extract all relevant information about {section.lower()} from the quote document.")
            
            # Group fields by type within section
            text_fields = [f for f in fields if f[1].get("type") == "string"]
            checkbox_fields = [f for f in fields if f[1].get("type") == "boolean"]
            
            # Add field-type specific instructions if needed
            if checkbox_fields and len(checkbox_fields) > 5:
                checkbox_examples = ", ".join([f"'{f[1].get('description', f[0])}'" for f in checkbox_fields[:3]])
                prompt_parts.append(f"  - This section contains many checkbox fields (e.g., {checkbox_examples}). Only mark a checkbox 'YES' if there is explicit evidence in the quote that the feature is included.")
            
            if text_fields and len(text_fields) > 5:
                text_examples = ", ".join([f"'{f[1].get('description', f[0])}'" for f in text_fields[:3]])
                prompt_parts.append(f"  - This section contains text fields (e.g., {text_examples}) that need specific values extracted from the quote.")
            
            # Add special instructions for specific sections
            if "Control & Programming" in section:
                prompt_parts.append("  - When determining PLC type, look for keywords like 'Allen Bradley', 'CompactLogix', 'B&R', or 'ControlLogix'.")
                prompt_parts.append("  - For HMI size, common options are 5.7\", 10\", and 15\". The size is often specified directly with the HMI.")
                prompt_parts.append("  - For beacon lights, if a 'tri-color beacon' or 'three-color tower light' is mentioned, mark all three colors (Red, Green, Yellow) as 'YES'.")
            
            if "Filling" in section:
                prompt_parts.append("  - For pump volumes, look for specific capacity values like '50cc', '100cc', '250cc', etc.")
                prompt_parts.append("  - Nozzle sizes should be compatible with the container being filled. Typical size is 2-3mm smaller than the container opening.")
            
            if "Reject" in section:
                prompt_parts.append("  - Mark all the specific reject reasons mentioned in the quote. Common ones include cap presence, fill weight, and label position.")
                prompt_parts.append("  - If a vision system is mentioned for inspection, look for specific inspection parameters it will check.")
            
            if "Labeling" in section:
                prompt_parts.append("  - For label application type, determine if it's a wrap-around, front/back, or multi-panel system.")
                prompt_parts.append("  - Check if a specific label head model (like LS100 or LS200) is mentioned.")
    
    # Add some general packaging industry interpretation guidelines
    prompt_parts.append("\n## PACKAGING INDUSTRY INTERPRETATION GUIDELINES:")
    prompt_parts.append("1. When a line item includes the phrase 'Including:' followed by features, mark all those features as 'YES' in the template.")
    prompt_parts.append("2. If a feature is described as 'Standard' or 'STD', it should be marked as 'YES'.")
    prompt_parts.append("3. For HMI and PLC types, carefully check for brand names in the machine descriptions.")
    prompt_parts.append("4. When filling systems mention 'Bottom-up filling' or 'Diving nozzles', these are specialized fill methods.")
    prompt_parts.append("5. If a multi-color beacon light is mentioned without specifying colors, mark all standard colors (Red, Yellow, Green) as 'YES'.")
    prompt_parts.append("6. For servo-controlled systems, look for associated parameters like torque range or speed.")
    
    return prompt_parts

def enhance_placeholder_context_with_outline(context_map: Dict[str, str], explicit_placeholder_mappings: Dict[str, str], outline_path: str = "full_fields_outline.md") -> Dict[str, str]:
    """
    Enhances the extracted placeholder context by cross-referencing with the structured outline file.
    
    Args:
        context_map: The existing context map from extract_placeholder_context_hierarchical
        explicit_placeholder_mappings: The dictionary of explicit mappings for the template.
        outline_path: Path to the full_fields_outline.md file
        
    Returns:
        Enhanced context map with more accurate hierarchical information
    """
    print(f"Enhancing placeholder context using outline file: {outline_path}")
    
    try:
        if not os.path.exists(outline_path):
            print(f"Warning: Outline file not found at {outline_path}")
            return context_map
            
        # Load the outline file
        with open(outline_path, 'r', encoding='utf-8') as f:
            outline_lines = f.readlines()
            
        # Build section mapping for common section names to standardized section names
        section_mapping = {
            # Control & Programming
            "control": "Control & Programming Specifications",
            "programming": "Control & Programming Specifications",
            "plc": "Control & Programming Specifications",
            "hmi": "Control & Programming Specifications",
            "batch": "Control & Programming Specifications",
            "beacon": "Control & Programming Specifications",
            "panel": "Control & Programming Specifications",
            "remote": "Control & Programming Specifications",
            "e-stop": "Control & Programming Specifications",
            "explosion proof": "Control & Programming Specifications",
            
            # Utility
            "utility": "Utility Specifications",
            "voltage": "Utility Specifications",
            "certification": "Utility Specifications",
            "conformity": "Utility Specifications",
            "hz": "Utility Specifications",
            "psi": "Utility Specifications",
            "amps": "Utility Specifications",
            "destination": "Utility Specifications",
            
            # Bottle Handling
            "bottle handling": "Bottle Handling System Specifications",
            "bottle": "Bottle Handling System Specifications",
            "vial": "Bottle Handling System Specifications",
            "puck": "Bottle Handling System Specifications",
            "container": "Bottle Handling System Specifications",
            "turn table": "Bottle Handling System Specifications",
            "index": "Bottle Handling System Specifications",
            "motion": "Bottle Handling System Specifications",
            "infeed": "Bottle Handling System Specifications",
            "outfeed": "Bottle Handling System Specifications",
            "transfer": "Bottle Handling System Specifications",
            "feeding": "Bottle Handling System Specifications",
            "elevator": "Bottle Handling System Specifications",
            "tube": "Bottle Handling System Specifications",
            "turn": "Bottle Handling System Specifications",
            "bulk": "Bottle Handling System Specifications",
            
            # Reject/Inspection
            "reject": "Reject / Inspection System",
            "inspection": "Reject / Inspection System",
            "verification": "Reject / Inspection System",
            
            # Liquid Filling
            "liquid filling": "Liquid Filling System Specifications",
            "filling": "Liquid Filling System Specifications",
            "pump": "Liquid Filling System Specifications",
            "nozzle": "Liquid Filling System Specifications",
            "volume": "Liquid Filling System Specifications",
            "valve": "Liquid Filling System Specifications",
            "tank": "Liquid Filling System Specifications",
            "hopper": "Liquid Filling System Specifications",
            "gutter": "Liquid Filling System Specifications",
            "cleaning": "Liquid Filling System Specifications",
            "weighing": "Liquid Filling System Specifications",
            "location": "Liquid Filling System Specifications",
            
            # Material
            "material": "Material Specifications",
            "product contact": "Material Specifications",
            "material certification": "Material Specifications",
            "tooling": "Material Specifications",
            
            # Parts & Materials
            "change part": "Change Part Quantities and Construction Materials",
            "construction material": "Change Part Quantities and Construction Materials",
            "seal": "Change Part Quantities and Construction Materials",
            "tubing": "Change Part Quantities and Construction Materials",
            "slat": "Change Part Quantities and Construction Materials",
            
            # Capping
            "capping": "Capping System Specifications",
            "cap": "Capping System Specifications",
            "torque": "Capping System Specifications",
            "placement": "Capping System Specifications",
            "sorting": "Capping System Specifications",
            "centering": "Capping System Specifications",
            "bulk feeder": "Capping System Specifications",
            
            # BeltStar
            "beltstar": "BeltStar System Specifications",
            "belt star": "BeltStar System Specifications",
            "belt driven": "BeltStar System Specifications",
            
            # Labeling
            "labeling": "Labeling System Specifications",
            "label": "Labeling System Specifications",
            "reel": "Labeling System Specifications",
            "application": "Labeling System Specifications",
            "separator": "Labeling System Specifications",
            
            # Coding
            "coding": "Coding and Inspection System Specifications",
            "coder": "Coding and Inspection System Specifications",
            "vision": "Coding and Inspection System Specifications",
            "print": "Coding and Inspection System Specifications",
            "ocr": "Coding and Inspection System Specifications",
            "ocv": "Coding and Inspection System Specifications",
            "barcode": "Coding and Inspection System Specifications",
            "videojet": "Coding and Inspection System Specifications",
            "laser": "Coding and Inspection System Specifications",
            
            # Induction
            "induction": "Induction Specifications",
            "enercon": "Induction Specifications",
            "sealing": "Induction Specifications",
            "tunnel": "Induction Specifications",
            
            # Shrink Sleeve
            "shrink sleeve": "Shrink Sleeve Specifications",
            "sleeve": "Shrink Sleeve Specifications",
            "shrink tunnel": "Shrink Sleeve Specifications",
            
            # Conveyor
            "conveyor": "Conveyor Specifications",
            "transport": "Conveyor Specifications",
            "chain": "Conveyor Specifications",
            
            # Others
            "gas purge": "Gas Purge",
            "nitrogen": "Gas Purge",
            "desiccant": "Desiccant",
            "cannister": "Desiccant",
            "roll / pouch": "Desiccant",
            "cottoner": "Cottoner",
            "cotton": "Cottoner",
            "plugging": "Plugging System Specifications",
            "plug": "Plugging System Specifications",
            "euro guarding": "Euro Guarding",
            "guarding": "Euro Guarding",
            "panel material": "Euro Guarding",
            "validation": "Validation Documents",
            "fat": "Validation Documents",
            "sat": "Validation Documents",
            "protocol": "Validation Documents",
            "acceptance test": "Validation Documents",
            "customer review": "Validation Documents",
            "documentation": "Validation Documents",
            "factory acceptance": "Validation Documents",
            "site acceptance": "Validation Documents",
            "test protocol": "Validation Documents",
            "support": "Validation Documents",
            "commissioning": "Validation Documents",
            "start-up": "Validation Documents",
            "manual": "Manual Specifications",
            "language": "Manual Specifications",
            "warranty": "Warranty & Install & Spares",
            "install": "Warranty & Install & Spares",
            "spares": "Warranty & Install & Spares",
            "packaging": "Packaging & Transport",
            "transport": "Packaging & Transport",
            "tablet counter": "Street Fighter Tablet Counter",
            "street fighter": "Street Fighter Tablet Counter",
            
            # Missing sections found in analysis
            "top hold down": "Labeling System Specifications",
            "top_hold_down": "Labeling System Specifications",
            "thd": "Labeling System Specifications",
            "lift fighter": "Street Fighter Tablet Counter",
            "lift_fighter": "Street Fighter Tablet Counter", 
            "lf": "Street Fighter Tablet Counter",
            "arm type": "Labeling System Specifications",
            "arm_type": "Labeling System Specifications",
            "switch type": "Euro Guarding",
            "switch_type": "Euro Guarding",
            "top cover": "Euro Guarding",
            "top_cover": "Euro Guarding",
            "no. of funnels": "Street Fighter Tablet Counter",
            "no._of_funnels": "Street Fighter Tablet Counter",
            "nf": "Street Fighter Tablet Counter",
            "options": "Liquid Filling System Specifications",
            
            # Additional mappings for remaining placeholders
            "type": "General",
            "argon": "Gas Purge",
            "linear": "Labeling System Specifications",
            "inline": "Labeling System Specifications",
            "rotary": "Labeling System Specifications",
            "none": "General",
            "unknown": "General",
            "other": "General",
            "found": "General"
        }
        
        # Special case handling for truncated entries
        truncated_mappings = {
            "bulk": "Bottle Handling System Specifications",
            "turn": "Bottle Handling System Specifications",
            "location": "Liquid Filling System Specifications"
        }
        
        # Parse the outline structure
        outline_context = {}
        section_hierarchy = {}  # Maps section names to their hierarchy for easier reference
        current_section = ""
        current_subsection = ""
        current_sub_subsection = ""
        
        # First pass: Extract section hierarchy
        for line in outline_lines:
            line = line.strip()
            if not line:
                continue
                
            # Main sections (## headings)
            if line.startswith('## '):
                current_section = line[3:].strip()
                current_subsection = ""
                current_sub_subsection = ""
                section_hierarchy[current_section.lower()] = {"name": current_section, "subsections": {}}
                
                # Add the main section to the outline context as well
                section_key = current_section.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace('&', 'and')
                outline_context[section_key] = current_section
            
            # Subsections (- headings with no indentation)
            elif line.startswith('- ') and not line.startswith('  - '):
                # Extract the subsection name (remove any (text), (section), etc.)
                subsection_text = line[2:].strip()
                clean_subsection = subsection_text
                if '(' in subsection_text:
                    clean_subsection = subsection_text.split('(')[0].strip()
                current_subsection = clean_subsection
                current_sub_subsection = ""
                
                if current_section.lower() in section_hierarchy:
                    section_hierarchy[current_section.lower()]["subsections"][current_subsection.lower()] = {
                        "name": current_subsection,
                        "items": {}
                    }
                
                # Special case - create field entries for checkbox subsections
                if '(checkbox)' in line:
                    field_key = clean_subsection.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace('&', 'and')
                    field_key = field_key + "_check"
                    
                    context_str = f"{current_section} - {clean_subsection}"
                    outline_context[field_key] = context_str
                
                # Add all subsections, regardless of type
                subsection_key = clean_subsection.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace('&', 'and')
                outline_context[subsection_key] = f"{current_section} - {clean_subsection}"
            
            # Sub-subsections (indented items with checkbox)
            elif line.startswith('  - '):
                # This is a field (checkbox or text)
                field_text = line[4:].strip()
                clean_field = field_text
                if '(' in field_text:
                    clean_field = field_text.split('(')[0].strip()
                
                # Store in hierarchy
                if current_section.lower() in section_hierarchy and current_subsection.lower() in section_hierarchy[current_section.lower()]["subsections"]:
                    section_hierarchy[current_section.lower()]["subsections"][current_subsection.lower()]["items"][clean_field.lower()] = clean_field
                
                # Create field key (for both checkbox and text fields)
                field_key = clean_field.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace('&', 'and').replace('"', '').replace('\'', '')
                
                # Add _check suffix for checkbox fields
                field_key_check = field_key + "_check"
                
                # Build context
                if current_subsection:
                    context_str = f"{current_section} - {current_subsection} - {clean_field}"
                else:
                    context_str = f"{current_section} - {clean_field}"
                
                if '(checkbox)' in line:
                    outline_context[field_key_check] = context_str
                    
                    # Also add version without _check for reference
                    outline_context[field_key] = context_str
            
                # For text fields
                elif '(text)' in line or '(qty)' in line:
                    # Build context
                    if current_subsection:
                        context_str = f"{current_section} - {current_subsection} - {clean_field}"
                    else:
                        context_str = f"{current_section} - {clean_field}"
                    
                    outline_context[field_key] = context_str
                    
                    # For any field that might be selectable, add a checkbox version too
                    if not any(term in line.lower() for term in ["comments", "file name", "capacity", "model", "example", "actual value"]):
                        outline_context[f"{field_key}_check"] = context_str
                    
            # For other types of fields (section, sub-section)
            elif line and not line.startswith('#'):
                # Skip this section if we're not in a valid section or if line is empty
                if not current_section or not line.strip():
                    continue
                    
                # For other types of entries, use the line text as field
                other_field_text = line.strip()
                clean_other_field = other_field_text
                if '(' in other_field_text:
                    clean_other_field = other_field_text.split('(')[0].strip()
                
                # Create field key
                other_field_key = clean_other_field.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace('&', 'and').replace('"', '').replace('\'', '')
                
                # Build context
                if current_subsection:
                    context_str = f"{current_section} - {current_subsection} - {clean_other_field}"
                else:
                    context_str = f"{current_section} - {clean_other_field}"
                
                outline_context[other_field_key] = context_str
                
                # For fields that might be selectable
                if not any(term in line.lower() for term in ["qty", "file name", "capacity", "model", "comments"]):
                    outline_context[f"{other_field_key}_check"] = context_str
        
        print(f"Parsed {len(outline_context)} fields from outline file")
        
        # Build a comprehensive mapping of placeholders based on the full_fields_outline.md
        # EXPANDED: Add even more well-known fields based on the template content
        known_fields = {
            # PLC types
            "plc_b&r_check": "Control & Programming Specifications - PLC - B & R",
            "plc_bandr_check": "Control & Programming Specifications - PLC - B & R",
            "plc_allen_bradley_check": "Control & Programming Specifications - PLC - Allen Bradley",
            "plc_allenb_check": "Control & Programming Specifications - PLC - Allen Bradley",
            "plc_compactlogix_check": "Control & Programming Specifications - PLC - CompactLogix",
            "plc_compactl_check": "Control & Programming Specifications - PLC - CompactLogix",
            "plc_controllogix_check": "Control & Programming Specifications - PLC - ControlLogix",
            "plc_controll_check": "Control & Programming Specifications - PLC - ControlLogix",
            
            # HMI sizes and types
            "hmi_10_check": "Control & Programming Specifications - HMI - Size - 10\"",
            "hmi_15_check": "Control & Programming Specifications - HMI - Size - 15\"",
            "hmi_5_7_check": "Control & Programming Specifications - HMI - Size - 5.7\" n/a for vision",
            "hmi_size10_check": "Control & Programming Specifications - HMI - Size - 10\"",
            "hmi_size15_check": "Control & Programming Specifications - HMI - Size - 15\"",
            "hmi_size5.7_check": "Control & Programming Specifications - HMI - Size - 5.7\" n/a for vision",
            "etr_10hmi_check": "Control & Programming Specifications - HMI - Size - 10\"",
            "hmi_pv10_check": "Control & Programming Specifications - HMI - Size - 10\"",
            "hmi_pv7_check": "Control & Programming Specifications - HMI - Size - 7\"",
            "hmi_allenb_check": "Control & Programming Specifications - HMI - Allen Bradley",
            "hmi_b&r_check": "Control & Programming Specifications - HMI - B & R",
            "hmi_pc_check": "Control & Programming Specifications - HMI - PC Upgrade",
            "lan_e_check": "Control & Programming Specifications - HMI - Language - English",
            "lan_f_check": "Control & Programming Specifications - HMI - Language - French",
            
            # Other control/programming specs
            "explosion_proof_check": "Control & Programming Specifications - Explosion proof",
            "cpp_1axis_check": "Control & Programming Specifications - Control Panel Post - 1 Axis",
            "cpp_2axis_check": "Control & Programming Specifications - Control Panel Post - 2 Axis",
            "cpp_3axis_check": "Control & Programming Specifications - Control Panel Post - 3 Axis",
            "cpp_fixed_check": "Control & Programming Specifications - Control Panel Post - Fixed (STD for Explosive Environment)",
            
            # Batch reporting and data
            "batch_none_check": "Control & Programming Specifications - Batch / Data Report - None",
            "batch_yes15_check": "Control & Programming Specifications - Batch / Data Report - Yes (requires 15\" HMI)",
            "batch_at_check": "Control & Programming Specifications - Batch / Data Report - Audit Trail",
            "batch_sht_check": "Control & Programming Specifications - Batch / Data Report - Summary Header with Tracking",
            
            # Beacon lights
            "blt_red_check": "Control & Programming Specifications - Beacon Light Tower - Red",
            "blt_green_check": "Control & Programming Specifications - Beacon Light Tower - Green",
            "blt_yellow_check": "Control & Programming Specifications - Beacon Light Tower - Yellow",
            "blt_audible_check": "Control & Programming Specifications - Beacon Light Tower - Audible",
            "blt_none_check": "Control & Programming Specifications - Beacon Light Tower - None",
            
            # Utility specifications
            "voltage": "Utility Specifications - Voltage",
            "hz": "Utility Specifications - Hz",
            "psi": "Utility Specifications - PSI",
            "amps": "Utility Specifications - AMPS", 
            "cfm": "Utility Specifications - CFM",
            "ce_csa_check": "Utility Specifications - Certification - CSA",
            "ce_ce_check": "Utility Specifications - Certification - CE",
            "ce_expl_check": "Utility Specifications - Certification - Explosion",
            "ce_c1d2_check": "Utility Specifications - Certification - Class1 Div 2",
            "ce_none_check": "Utility Specifications - Certification - None",
            
            # Reject/inspection system
            "cap_prs_check": "Reject / Inspection System - Reject Reasons - Cap Prs.",
            "rr_fillwt_check": "Reject / Inspection System - Reject Reasons - Fill WT / Count",
            "rr_nocot_check": "Reject / Inspection System - Reject Reasons - No Cotton",
            "rr_stempr_check": "Reject / Inspection System - Reject Reasons - Stem Prs.",
            "rr_pintpr_check": "Reject / Inspection System - Reject Reasons - Pintle Prs.",
            "rr_nocap_check": "Reject / Inspection System - Reject Reasons - Cap Prs.",
            "rr_plugpr_check": "Reject / Inspection System - Reject Reasons - Plug Prs.",
            "rr_hxthrd_check": "Reject / Inspection System - Reject Reasons - High/Cross THRD",
            "rr_nofoil_check": "Reject / Inspection System - Reject Reasons - No Foil",
            "rr_torque_check": "Reject / Inspection System - Reject Reasons - Torque",
            "rr_ocv_check": "Reject / Inspection System - Reject Reasons - OCV",
            "rr_ocr_check": "Reject / Inspection System - Reject Reasons - OCR",
            "rr_barcode_check": "Reject / Inspection System - Reject Reasons - Bar Code",
            "rr_labpos_check": "Reject / Inspection System - Reject Reasons - Label Position",
            "rr_labprs_check": "Reject / Inspection System - Reject Reasons - Label Prs",
            "rr_nodesc_check": "Reject / Inspection System - Reject Reasons - No Desiccant",
            
            # Material specifications
            "cap_ss304_check": "Material Specifications - Cap - SS 304",
            "cap_ss316_check": "Material Specifications - Cap - SS 316",
            "cap_uins_check": "Material Specifications - Cap - Urethane Insert",
            "bot_hdpe_check": "Material Specifications - Bottle - HDPE",
            "bot_antist_check": "Material Specifications - Bottle - Antistatic",
            
            # Bottle handling
            "bs_cpbo_check": "BeltStar System Specifications - Cap Placement - Belts – on the fly",
            "bs_cpnone_check": "BeltStar System Specifications - Cap Placement - None",
            "bs_cpac_check": "BeltStar System Specifications - Cap Placement - AC",
            "bs_tam_check": "BeltStar System Specifications - Torque - Air motor",
            "bs_tmc_check": "BeltStar System Specifications - Torque - Magnet clutch",
            "bs_ts_check": "BeltStar System Specifications - Torque - Servo",
            "bs_tb_check": "BeltStar System Specifications - Torque - Belts",
            "bs_tadb_check": "BeltStar System Specifications - Torque - AC motor DC Brake",
            "bs_tf_check": "BeltStar System Specifications - Torque - Feedback",
            "bs_tht_check": "BeltStar System Specifications - Torque - HMI Adj. Torque",
            "bs_csds_check": "BeltStar System Specifications - Cap Sorting - Docking Station",
            "bs_cse_check": "BeltStar System Specifications - Cap Sorting - Elevator",
            "bs_manone_check": "BeltStar System Specifications - Motorized Adj. - None",
            "bs_may_check": "BeltStar System Specifications - Motorized Adj. - Yes",
            
            # Cottoner system
            "c_none_check": "Cottoner - None",
            "c_sp_check": "Cottoner - Sensing - Presence",
            "c_cn_check": "Cottoner - Cotton Bin - No",
            "c_cy_check": "Cottoner - Cotton Bin - Yes",
            
            # Validation and documentation
            "vd_d_check": "Validation Documents - Design Qualification (DQ) Documentation",
            "vd_f_check": "Validation Documents - Factory Acceptance Test (FAT) Protocol Package - Includes documentation, support, and customer review",
            "vd_fd_check": "Validation Documents - Functional/Design Specification (FS/DS)",
            "vd_h_check": "Validation Documents - Hardware/Software Design Specification (HDS/SDS)",
            "vd_i_check": "Validation Documents - Installation/Operational Qualification (IQ/OQ)",
            "vd_s_check": "Validation Documents - Site Acceptance Test (SAT) Protocol Package - Includes documentation, support, and customer review",
            
            # Guarding
            "eg_pnl_check": "Euro Guarding - Panel material - Lexan",
            "eg_pmtg_check": "Euro Guarding - Panel material - Tempered glass",
            
            # Capping system
            "cs_csm_check": "Capping System Specifications - Cap Sorting - Mechanical",
            
            # Plugging system
            "ps_plm_check": "Plugging System Specifications - Plug Placement - Mechanical",
            "ps_psm_check": "Plugging System Specifications - Plug Sorting - Mechanical"
        }
        
        # Add all known fields to the outline context
        for field, context in known_fields.items():
            outline_context[field] = context
        
        # Create normalized versions of context map keys for better matching
        normalized_context_keys = {}
        for key in context_map.keys():
            # Create normalized versions
            norm_key1 = key.lower()  # lowercase
            norm_key2 = key.lower().replace('_', '')  # no underscores
            norm_key3 = key.lower().replace('_check', '')  # without _check suffix
            norm_key4 = key.lower().replace('_check', '').replace('_', '')  # clean
            
            # Additional normalizations for special cases
            norm_key5 = re.sub(r'[^a-z0-9]', '', key.lower())  # alphanumeric only
            norm_key6 = re.sub(r'[0-9]+"?', '', key.lower())  # remove numbers and inch marks
            
            normalized_context_keys[norm_key1] = key
            normalized_context_keys[norm_key2] = key
            normalized_context_keys[norm_key3] = key
            normalized_context_keys[norm_key4] = key
            normalized_context_keys[norm_key5] = key
            normalized_context_keys[norm_key6] = key
        
        # Create normalized versions of outline keys for better matching
        normalized_outline_keys = {}
        for key, value in outline_context.items():
            # Create multiple normalized versions for more robust matching
            norm_key1 = key.lower().replace('_', '')  # No underscores
            norm_key2 = key.lower()  # With underscores
            norm_key3 = key.lower().replace('_check', '')  # Without _check suffix
            norm_key4 = key.lower().replace('_check', '').replace('_', '')  # Without _check and no underscores
            norm_key5 = re.sub(r'[^a-z0-9]', '', key.lower())  # alphanumeric only
            norm_key6 = re.sub(r'[0-9]+"?', '', key.lower())  # remove numbers and inch marks
            
            normalized_outline_keys[norm_key1] = key
            normalized_outline_keys[norm_key2] = key
            normalized_outline_keys[norm_key3] = key
            normalized_outline_keys[norm_key4] = key
            normalized_outline_keys[norm_key5] = key
            normalized_outline_keys[norm_key6] = key
            
            # Handle common abbreviations and expanded forms
            if 'plc' in key.lower():
                expanded = key.lower().replace('plc', 'programmable_logic_controller')
                normalized_outline_keys[expanded] = key
            if 'hmi' in key.lower():
                expanded = key.lower().replace('hmi', 'human_machine_interface')
                normalized_outline_keys[expanded] = key
            if 'vfd' in key.lower():
                expanded = key.lower().replace('vfd', 'variable_frequency_drive')
                normalized_outline_keys[expanded] = key
            if 'ss' in key.lower():
                expanded = key.lower().replace('ss', 'stainless_steel')
                normalized_outline_keys[expanded] = key
            if 'qty' in key.lower():
                expanded = key.lower().replace('qty', 'quantity')
                normalized_outline_keys[expanded] = key
            if 'temp' in key.lower():
                expanded = key.lower().replace('temp', 'temperature')
                normalized_outline_keys[expanded] = key
                
            # Add common alternative terms
            if 'bottle' in key.lower():
                alt = key.lower().replace('bottle', 'container')
                normalized_outline_keys[alt] = key
            if 'container' in key.lower():
                alt = key.lower().replace('container', 'bottle')
                normalized_outline_keys[alt] = key
            if 'plug' in key.lower():
                alt = key.lower().replace('plug', 'stopper')
                normalized_outline_keys[alt] = key
            if 'cap' in key.lower():
                alt = key.lower().replace('cap', 'lid')
                normalized_outline_keys[alt] = key
            
            # Handle size notation variations
            size_patterns = [
                (r'(\d+)_?inch', r'\1in'),
                (r'(\d+)in', r'\1"'),
                (r'(\d+)"', r'\1inch')
            ]
            for pattern, replacement in size_patterns:
                if re.search(pattern, key.lower()):
                    alt = re.sub(pattern, replacement, key.lower())
                    normalized_outline_keys[alt] = key
        
        print(f"Created {len(normalized_outline_keys)} normalized outline keys")
        
        # Enhance the context map with outline information
        enhanced_context_map = context_map.copy()
        enhanced_count = 0
        
        # Print some debug info about existing keys (reduced verbosity)
        print(f"Original context map has {len(context_map)} keys")
        
        # First pass: Try direct key matches and known fields
        for key in context_map.keys():
            # First, check if this is a specific placeholder with predefined context
            if key in explicit_placeholder_mappings:
                enhanced_context_map[key] = explicit_placeholder_mappings[key]
                enhanced_count += 1
                continue
                
            # First, check if this is a known field with predefined context
            if key in known_fields:
                enhanced_context_map[key] = known_fields[key]
                enhanced_count += 1
                continue
                
            # Try direct match next
            if key in outline_context:
                enhanced_context_map[key] = outline_context[key]
                enhanced_count += 1
                continue
                
            # Try normalized keys
            norm_key1 = key.lower().replace('_', '')
            norm_key2 = key.lower()
            norm_key3 = key.lower().replace('_check', '')
            norm_key4 = key.lower().replace('_check', '').replace('_', '')
            
            # Try each normalized version
            for norm_key in [norm_key1, norm_key2, norm_key3, norm_key4]:
                if norm_key in normalized_outline_keys:
                    outline_key = normalized_outline_keys[norm_key]
                enhanced_context_map[key] = outline_context[outline_key]
                enhanced_count += 1
                break
            
            # If we enhanced the context, continue to next key
            if enhanced_context_map[key] != context_map[key]:
                continue
                
            # Special case for PLC brands
            if 'plc' in key.lower() and 'check' in key.lower():
                if 'b&r' in key.lower() or 'br' in key.lower() or 'bandr' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - PLC - B & R"
                    enhanced_count += 1
                    continue
                if 'allen' in key.lower() or 'bradley' in key.lower() or 'ab' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - PLC - Allen Bradley"
                    enhanced_count += 1
                    continue
                if 'compact' in key.lower() or 'compactl' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - PLC - CompactLogix"
                    enhanced_count += 1
                    continue
                if 'control' in key.lower() or 'controll' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - PLC - ControlLogix"
                    enhanced_count += 1
                    continue
            
            # Special case for HMI sizes
            if 'hmi' in key.lower() and 'check' in key.lower():
                if '10' in key.lower() or '10inch' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - HMI - Size - 10\""
                    enhanced_count += 1
                    continue
                if '15' in key.lower() or '15inch' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - HMI - Size - 15\""
                    enhanced_count += 1
                    continue
                if '5.7' in key.lower() or '5_7' in key.lower() or '57' in key.lower():
                    enhanced_context_map[key] = "Control & Programming Specifications - HMI - Size - 5.7\" n/a for vision"
                    enhanced_count += 1
                    continue
            
            # Try using the section mapping for remaining unmatched keys
            if enhanced_context_map[key] == context_map[key]:  # Not enhanced yet
                original_context = context_map[key].lower()
                
                # Check for truncated context
                for truncated_term, section in truncated_mappings.items():
                    if original_context.startswith(truncated_term):
                        # Use the full context but replace the section
                        enhanced_context_map[key] = f"{section} - {context_map[key]}"
                        enhanced_count += 1
                        continue
                
                # If still not enhanced, look for section keywords
                if enhanced_context_map[key] == context_map[key]:
                    # Look for section keywords in the original context
                    detected_section = None
                    detected_subsection = None
                    max_match_score = 0
                    
                    # First pass: try to find a direct section match
                    for keyword, section in section_mapping.items():
                        if keyword in original_context:
                            # Basic match score - longer keyword matches are more specific
                            match_score = len(keyword)
                            if match_score > max_match_score:
                                max_match_score = match_score
                                detected_section = section
                
                if detected_section:
                    # Look for subsection matches within this section
                    if detected_section in section_hierarchy:
                        section_data = section_hierarchy[detected_section.lower()]
                        for subsection_key, subsection_data in section_data.get("subsections", {}).items():
                            if subsection_key in original_context.lower():
                                detected_subsection = subsection_data.get("name", "")
                                break
                    
                    # Create an improved context using the detected section and subsection
                    if detected_subsection:
                        # Keep any additional context beyond the subsection
                        parts = original_context.split(' - ')
                        field_part = parts[-1] if len(parts) > 1 and parts[-1] not in detected_section.lower() and parts[-1] not in detected_subsection.lower() else key
                        enhanced_context_map[key] = f"{detected_section} - {detected_subsection} - {field_part}"
                    else:
                        # Just use the section with the original field
                        parts = original_context.split(' - ')
                        field_part = parts[-1] if len(parts) > 1 and parts[-1] not in detected_section.lower() else key
                        enhanced_context_map[key] = f"{detected_section} - {field_part}"
                    enhanced_count += 1
                
                # Second pass: try extracting information from the key itself if not yet matched
                if enhanced_context_map[key] == context_map[key]:  # Still not enhanced
                    key_lower = key.lower()
                    
                    # Check for section indicators in the key
                    for keyword, section in section_mapping.items():
                        if keyword in key_lower:
                            # Create an improved context based on the key
                            enhanced_context_map[key] = f"{section} - {key}"
                            enhanced_count += 1
                            break
                
                # Third pass: check for known patterns in keys
                if enhanced_context_map[key] == context_map[key]:  # Still not enhanced
                    key_lower = key.lower()
                    
                    # PLC related fields
                    if 'plc' in key_lower:
                        enhanced_context_map[key] = "Control & Programming Specifications - PLC"
                        enhanced_count += 1
                    
                    # HMI related fields
                    elif 'hmi' in key_lower:
                        enhanced_context_map[key] = "Control & Programming Specifications - HMI"
                        enhanced_count += 1
                    
                    # Utility related fields
                    elif any(term in key_lower for term in ['volt', 'hz', 'psi', 'amp', 'cfm']):
                        enhanced_context_map[key] = "Utility Specifications"
                        enhanced_count += 1
                    
                    # Bottle handling related fields
                    elif any(term in key_lower for term in ['bottle', 'container', 'vial', 'puck']):
                        enhanced_context_map[key] = "Bottle Handling System Specifications"
                        enhanced_count += 1
                    
                    # Reject system related fields
                    elif any(term in key_lower for term in ['reject', 'inspection', 'verification']):
                        enhanced_context_map[key] = "Reject / Inspection System"
                        enhanced_count += 1
                    
                    # Cap related fields
                    elif any(term in key_lower for term in ['cap', 'torque', 'capping']):
                        enhanced_context_map[key] = "Capping System Specifications"
                        enhanced_count += 1
                    
                    # Filling related fields
                    elif any(term in key_lower for term in ['fill', 'pump', 'nozzle']):
                        enhanced_context_map[key] = "Liquid Filling System Specifications"
                enhanced_count += 1
        
        # Print some statistics
        print(f"Enhanced {enhanced_count} of {len(context_map)} placeholder contexts")
        
        return enhanced_context_map
    
    except Exception as e:
        print(f"Error enhancing placeholder context: {e}")
        import traceback
        traceback.print_exc()
        return context_map

def parse_full_fields_outline(outline_md_content: str) -> Dict[str, Dict[str, list]]:
    """
    Parses the full_fields_outline.md content to extract a hierarchical structure
    of sections and subsections.

    Args:
        outline_md_content: The string content of the full_fields_outline.md file.

    Returns:
        A dictionary where keys are section names, and values are dictionaries
        containing a list of subsection names under "_subsections_".
        Example: 
        {
            "Section Name 1": {
                "_subsections_": ["Subsection Name 1.1", "Subsection Name 1.2"],
                "_fields_": [] # To be populated later with field objects
            },
            "Section Name 2": {
                "_subsections_": [], # No subsections, fields will go into _fields_
                "_fields_": []
            }
        }
    """
    parsed_structure = {}
    current_section_name = None
    # current_subsection_name = None # Not needed here, just collecting subsection titles

    for line in outline_md_content.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith('## '):
            current_section_name = line[3:].strip()
            if current_section_name not in parsed_structure:
                # Initialize with keys for subsections and direct fields for this section
                parsed_structure[current_section_name] = {"_subsections_": [], "_fields_": []}
            # current_subsection_name = None # Reset for new section
        elif line.startswith('- ') and current_section_name:
            # This is potentially a subsection title or a field listed directly under a section.
            # We only capture it as a subsection if it does NOT look like a field specifier.
            item_name_full = line[2:].strip()
            
            # If it doesn't have (checkbox), (text), (qty), (section) it is likely a true subsection title
            if not re.search(r'\((checkbox|text|qty|section)\)', item_name_full, re.IGNORECASE):
                subsection_title = re.sub(r'\s*\(.*\)\s*$', '', item_name_full).strip() # Clean name
                if subsection_title and subsection_title not in parsed_structure[current_section_name]["_subsections_"]:
                    parsed_structure[current_section_name]["_subsections_"].append(subsection_title)
    
    # Sort subsections alphabetically within each section for consistent order
    for section_name in parsed_structure:
        if "_subsections_" in parsed_structure[section_name]:
            parsed_structure[section_name]["_subsections_"] = sorted(list(set(parsed_structure[section_name]["_subsections_"])))

    return parsed_structure

if __name__ == '__main__':
    test_template_path = 'template.docx' 
    if os.path.exists(test_template_path):
        print(f"Testing with template: {test_template_path}")
        
        placeholders = extract_placeholders(test_template_path)
        print("\n--- Unique Placeholders Found ---")
        if placeholders:
            for p_holder in placeholders:
                print(f"- {p_holder}")
        else:
            print("No placeholders found.")

        print("\n--- Placeholder Context (Hierarchical Test) ---")
        context = extract_placeholder_context_hierarchical(test_template_path, explicit_placeholder_mappings=DEFAULT_EXPLICIT_MAPPINGS)
        if context:
            count = 0
            for p_holder, desc in sorted(context.items()): 
                print(f"- '{p_holder}': '{desc}'")
                count += 1
                if count >= 50 and len(context) > 50: 
                    print("... and more.")
                    break
        else:
            print("No placeholder context generated.")
            
    else:
        print(f"Test template file not found at: {test_template_path}. Please provide a valid path.") 