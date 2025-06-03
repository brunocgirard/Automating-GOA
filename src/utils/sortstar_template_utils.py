import re
from typing import List, Set, Dict, Optional
from docx import Document
import os # Added for __main__ example

# Define explicit placeholder mappings for SortStar
explicit_placeholder_mappings = {
    "customer": "GENERAL ORDER ACKNOWLEDGEMENT > Customer",
    "machine": "GENERAL ORDER ACKNOWLEDGEMENT > Machine",
    "direction": "GENERAL ORDER ACKNOWLEDGEMENT > Direction",
    "quote": "Order Identification > Quote",
    "production_speed": "Order Identification > Production speed",
    "voltage": "Utility Specifications > Voltage",
    "phases": "Utility Specifications > Phases",
    "hz": "Utility Specifications > Hz",
    "amps": "Utility Specifications > AMPS",
    "psi": "Utility Specifications > PSI",
    "cfm": "Utility Specifications > CFM",
    "conformity_csa_check": "Utility Specifications > Conformity CSA Check",
    "ce_none_check": "Utility Specifications > CE None Check",
    "ce_csa_check": "Utility Specifications > CE CSA Check",
    "ce_expl_check": "Utility Specifications > CE Expl Check",
    "country": "Utility Specifications > Country of destination",
    "bs_984_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 220VAC 3 Phases LEFT TO RIGHT",
    "bs_1230_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 220VAC 3 Phases RIGHT TO LEFT",
    "bs_985_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 480VAC & 380VAC 3 Phases LEFT TO RIGHT",
    "bs_1229_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 18ft3 480VAC & 380VAC 3 Phases RIGHT TO LEFT",
    "bs_1264_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 24ft3 220VAC 3 Phases LEFT TO RIGHT",
    "bs_1265_check": "BASIC SYSTEMS > Mechanical Basic Machine configuration > Sortstar 24ft3 480VAC & 380VAC 3 Phases LEFT TO RIGHT",
    "op_2409_check": "OPTIONAL SYSTEMS > Guarding System > Lexan Euroguard Top Cover",
    "op_ias_check": "OPTIONAL SYSTEMS > Guarding System > Ionized Air System",
    "op_nrck_check": "OPTIONAL SYSTEMS > Guarding System > Requires Cap Kit",
    "el_0369_check": "OPTIONAL SYSTEMS > Electrical > Stack Light with Buzzer Option",
    "cps_ep_check": "OPTIONAL SYSTEMS > Control Specifications > Explosion proof",
    "cps_none_check": "OPTIONAL SYSTEMS > Control Specifications > None",
    "plc_b&r_check": "OPTIONAL SYSTEMS > Control Specifications > PLC B & R",
    "plc_compactl_check": "OPTIONAL SYSTEMS > Control Specifications > CompactLogix",
    "plc_controll_check": "OPTIONAL SYSTEMS > Control Specifications > ControlLogix",
    "hmi_b&r_check": "OPTIONAL SYSTEMS > Control Specifications > HMI B & R",
    "hmi_allenb_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Allen Bradley",
    "hmi_pc_check": "OPTIONAL SYSTEMS > Control Specifications > HMI PC Upgrade",
    "hmi_size5.7_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Size 5.7\"",
    "hmi_size10_check": "OPTIONAL SYSTEMS > Control Specifications > HMI Size 10\"",
    "cpp_1axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 1 Axis",
    "cpp_2axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 2 Axis – U-shaped",
    "cpp_3axis_check": "OPTIONAL SYSTEMS > Control Specifications > Control Panel Post 3 Axis",
    "rts_secomea_check": "OPTIONAL SYSTEMS > Remote Technical Service (Secomea)",
    "rts_none_check": "OPTIONAL SYSTEMS > Remote Technical Service > None",
    "rts_co_check": "OPTIONAL SYSTEMS > Remote Technical Service > Connection Only",
    "eg_none_check": "OPTIONAL SYSTEMS > Euro guarding > None",
    "eg_pnl_check": "OPTIONAL SYSTEMS > Euro guarding > Panel material Lexan",
    "eg_pmtg_check": "OPTIONAL SYSTEMS > Euro guarding > Tempered glass",
    "eg_stkw_check": "OPTIONAL SYSTEMS > Euro guarding > Switch type Key switch",
    "eg_stm_check": "OPTIONAL SYSTEMS > Euro guarding > Switch type Magnetic",
    "eg_sto_check": "OPTIONAL SYSTEMS > Euro guarding > Other",
    "eg_rc_check": "OPTIONAL SYSTEMS > Euro guarding > Reject Cover",
    "eg_rcnone_check": "OPTIONAL SYSTEMS > Euro guarding > Reject Cover None",
    "vd_f_check": "OPTIONAL SYSTEMS > Validation Documents > FAT",
    "vd_s_check": "OPTIONAL SYSTEMS > Validation Documents > SAT",
    "vd_d_check": "OPTIONAL SYSTEMS > Validation Documents > DQ",
    "vd_h_check": "OPTIONAL SYSTEMS > Validation Documents > HDS/SDS",
    "vd_fd_check": "OPTIONAL SYSTEMS > Validation Documents > FS/DS",
    "vd_i_check": "OPTIONAL SYSTEMS > Validation Documents > IQ/OQ",
    "pt_tcc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Transport charges Capmatic",
    "pt_tccu_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Transport charges Customer",
    "pt_pbc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging by Capmatic",
    "pt_pbcu_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging by Customer / Not Incl.",
    "pt_pts_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Skid",
    "pt_ptc_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Crate",
    "pt_ptsw_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Packaging type Sea Worthy",
    "wi_war1_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Warranty 1YR",
    "wi_war2_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Warranty 2YR",
    "sk_none_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit None",
    "sk_1yr_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit 1YR",
    "sk_2yr_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Spares Kit 2YR",
    # "rts_none_check" is duplicated, mapped to "OPTIONAL SYSTEMS > Remote Technical Service > None" above.
    # This second instance comes from: "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Remote Tech. Service None"
    # For now, I will keep the first mapping, but this might need review if context is critical.
    # To avoid key collision, I will temporarily comment out the second one or decide on a naming convention if necessary.
    # For now, let's assume the earlier "Remote Technical Service > None" is the primary one.
    # "rts_none_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Remote Tech. Service None",
    "rts_y_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Remote Tech. Service Yes",
    "stpc_none_check": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Start-up Commissioning None",
    "stpc_yes": "OPTIONAL SYSTEMS > Packaging & Transport & Warranty & Install & Spares > Start-up Commissioning Yes – no. of days"
}

def extract_placeholders(template_path: str) -> List[str]:
    """Reads the template and extracts all unique, cleaned placeholder keys."""
    print(f"Extracting placeholders from: {template_path}")
    placeholders: Set[str] = set()
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")

        for para in doc.paragraphs:
            for match in regex.findall(para.text):
                cleaned_key = match.strip()
                if cleaned_key:
                    placeholders.add(cleaned_key)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in regex.findall(cell.text):
                        cleaned_key = match.strip()
                        if cleaned_key:
                            placeholders.add(cleaned_key)
                            
        if not placeholders:
             print(f"Warning: No placeholders found in {template_path}")
             return []

        print(f"Found {len(placeholders)} unique placeholders.")
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error reading placeholders from template '{template_path}': {e}")
        return []

def is_likely_section_header(paragraph) -> bool:
    """Heuristically determines if a paragraph is a section header."""
    text = paragraph.text.strip()
    if not text: # Empty paragraphs are not headers
        return False
    
    if text.isupper() and len(text.split()) < 7:
        return True
        
    if paragraph.runs and paragraph.runs[0].bold and len(text.split()) < 7:
        if sum(1 for char in text if char.islower()) < len(text) / 2:
             return True
    return False

def extract_placeholder_context_hierarchical(template_path: str, 
                                            enhance_with_outline: bool = True, # For Sortstar, we might not have an outline yet
                                            outline_path: str = "sortstar_fields_outline.md", # Placeholder for potential future outline
                                            check_if_all_mapped: bool = True) -> Dict[str, str]:
    """
    Parses the template to extract placeholders and attempts to build hierarchical context
    by identifying section headers. Optionally enhances with outline file.
    
    Args:
        template_path: Path to the Word document template
        enhance_with_outline: Whether to enhance context with outline file
        outline_path: Path to the outline file
        check_if_all_mapped: If True, checks if all placeholders are already in explicit_placeholder_mappings
                             and skips extraction if they are
    """
    print(f"Extracting hierarchical placeholder context from (SortStar): {template_path}")
    
    if check_if_all_mapped:
        all_placeholders = extract_placeholders(template_path)
        if all_placeholders: # Check if list is not empty
            all_mapped = all(ph in explicit_placeholder_mappings for ph in all_placeholders)
            if all_mapped:
                print(f"All {len(all_placeholders)} placeholders are explicitly mapped for SortStar. Using explicit mappings.")
                return {ph: explicit_placeholder_mappings[ph] for ph in all_placeholders if ph in explicit_placeholder_mappings} # Ensure key exists
            else:
                unmapped = [ph for ph in all_placeholders if ph not in explicit_placeholder_mappings]
                print(f"Found {len(unmapped)} unmapped placeholders out of {len(all_placeholders)} for SortStar. Proceeding with dynamic extraction for unmapped.")
                if len(unmapped) <= 10:
                    print(f"Unmapped SortStar placeholders: {', '.join(unmapped)}")
        elif not all_placeholders and os.path.exists(template_path): # No placeholders found, but file exists
             print(f"Warning: No placeholders extracted by extract_placeholders for {template_path}, but file exists. Will rely on explicit_placeholder_mappings if any match.")
             # Fallback to returning all explicit mappings if dynamic extraction yields nothing
             return explicit_placeholder_mappings.copy()


    context_map: Dict[str, str] = {}
    # Initialize with explicit mappings, dynamic extraction will update/add to these
    context_map.update(explicit_placeholder_mappings)

    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")
        
        current_section_header = "General" 
        current_subsection_header = ""

        placeholder_details = {}

        for p_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if is_likely_section_header(para):
                current_section_header = para_text.replace(":","").strip()
                current_subsection_header = "" 
                continue
            
            elif para.runs and para.runs[0].bold and len(para_text.split()) < 5 and not para_text.isupper() and para_text.endswith(":"):
                current_subsection_header = para_text.replace(":","").strip()
                continue

            for r_match in regex.finditer(para_text):
                ph_key = r_match.group(1).strip()
                if ph_key and ph_key not in explicit_placeholder_mappings: # Only process if not explicitly mapped
                    if ph_key not in placeholder_details:
                        preceding_text = para_text[:r_match.start()].strip()
                        preceding_text = regex.sub("", preceding_text).strip().replace(":","").strip()
                        placeholder_details[ph_key] = {
                            "immediate_label": preceding_text if preceding_text else ph_key,
                            "section": current_section_header,
                            "subsection": current_subsection_header,
                            "type": "paragraph"
                        }

        for t_idx, table in enumerate(doc.tables):
            current_table_group_label = "" 
            for r_idx, row in enumerate(table.rows):
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text and not regex.search(first_cell_text) and len(first_cell_text.split()) < 4:
                        if r_idx == 0 or (r_idx > 0 and table.cell(r_idx,0).text != table.cell(r_idx-1,0).text):
                            current_table_group_label = first_cell_text.replace(":","").strip()
                        elif not current_table_group_label:
                            current_table_group_label = first_cell_text.replace(":","").strip()
                
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for r_match in regex.finditer(cell_text):
                        ph_key = r_match.group(1).strip()
                        if ph_key and ph_key not in explicit_placeholder_mappings: # Only process if not explicitly mapped
                            if ph_key not in placeholder_details: 
                                immediate_label = ""
                                if c_idx > 0: 
                                    label_cell_text = row.cells[c_idx-1].text.strip().replace(":","").strip()
                                    if label_cell_text and not regex.search(label_cell_text):
                                        immediate_label = label_cell_text
                                if not immediate_label: 
                                    immediate_label = cell_text[:r_match.start()].strip().replace(":","").strip()
                                
                                placeholder_details[ph_key] = {
                                    "immediate_label": immediate_label if immediate_label else ph_key,
                                    "section": current_section_header,
                                    "subsection": current_subsection_header, 
                                    "table_group": current_table_group_label if current_table_group_label != immediate_label else "",
                                    "type": "table"
                                }

        for ph_key, details in placeholder_details.items():
            parts = []
            if details.get("section") and details["section"] != "General": parts.append(details["section"])
            if details.get("subsection"): parts.append(details["subsection"])
            if details.get("table_group"): parts.append(details["table_group"])
            if details.get("immediate_label") and details["immediate_label"] != ph_key: parts.append(details["immediate_label"])
            
            if parts:
                context_map[ph_key] = " - ".join(filter(None, parts))
            elif ph_key not in context_map : # Add only if not already from explicit or other dynamic step
                context_map[ph_key] = ph_key 

        all_phs_from_doc = extract_placeholders(template_path)
        for ph in all_phs_from_doc:
            if ph not in context_map: # If any placeholder from doc is still not in map
                context_map[ph] = ph # Add with key as context as a fallback

        if not context_map: print(f"Warning: No placeholder context generated for SortStar template {template_path}")
        else: print(f"Generated/updated hierarchical context for {len(context_map)} SortStar placeholders.")
        
        # For SortStar, outline enhancement might be simpler or use a different outline file
        if enhance_with_outline and os.path.exists(outline_path):
            # This function would need to be defined or adapted for SortStar if its outline differs significantly
            # For now, we assume a similar enhancement logic can be applied if an outline exists.
            # context_map = enhance_placeholder_context_with_sortstar_outline(context_map, outline_path)
            print(f"Outline enhancement for SortStar to be implemented if {outline_path} is used and structured.")
        
        return context_map

    except Exception as e:
        print(f"Error extracting hierarchical placeholder context for SortStar from '{template_path}': {e}")
        import traceback
        traceback.print_exc()
        return explicit_placeholder_mappings # Fallback to explicit if error

def extract_placeholder_schema(template_path: str) -> Dict[str, Dict]:
    """
    Creates a structured JSON schema from the SortStar template.
    This version prioritizes explicit_placeholder_mappings.
    """
    print(f"Extracting JSON schema from SortStar template: {template_path}")
    
    schema: Dict[str, Dict] = {}
    
    # Populate schema with explicitly mapped placeholders first
    for ph_key, description in explicit_placeholder_mappings.items():
        field_type = "boolean" if ph_key.endswith("_check") else "string"
        # Basic section/subsection parsing from description string
        parts = description.split(" > ")
        section = parts[0] if len(parts) > 0 else "General"
        subsection = parts[1] if len(parts) > 1 else ""
        
        schema[ph_key] = {
            "type": field_type,
            "section": section,
            "subsection": subsection,
            "description": description, # Use full path as description
            "location": "unknown" # Location is harder to determine from flat list
        }
        if field_type == "boolean":
            # Simplified synonym/indicator generation for explicitly mapped items
            synonyms = [description.lower(), ph_key.replace("_check","").replace("_"," ")]
            schema[ph_key]["synonyms"] = list(set(synonyms))
            schema[ph_key]["positive_indicators"] = [f"with {s}" for s in synonyms] + ["yes", "selected", description.lower()]


    # Dynamically extract other placeholders if any are not in explicit_placeholder_mappings
    # This part largely mirrors the original template_utils but will add to the existing schema
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")
        
        current_section_header = "General"
        current_subsection_header = ""
        
        # Use a temporary dict for dynamically found details to merge later
        placeholder_details_dynamic = {}

        for p_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if is_likely_section_header(para):
                current_section_header = para_text.replace(":","").strip()
                current_subsection_header = ""
                continue
            elif para.runs and para.runs[0].bold and len(para_text.split()) < 5 and not para_text.isupper() and para_text.endswith(":"):
                current_subsection_header = para_text.replace(":","").strip()
                continue

            for r_match in regex.finditer(para_text):
                ph_key = r_match.group(1).strip()
                if ph_key and ph_key not in schema: # Only process if not already in schema from explicit map
                    preceding_text = para_text[:r_match.start()].strip()
                    preceding_text = regex.sub("", preceding_text).strip().replace(":","").strip()
                    placeholder_details_dynamic[ph_key] = {
                        "immediate_label": preceding_text if preceding_text else ph_key,
                        "section": current_section_header,
                        "subsection": current_subsection_header,
                        "type": "paragraph"
                    }

        for t_idx, table in enumerate(doc.tables):
            current_table_group_label = ""
            for r_idx, row in enumerate(table.rows):
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text and not regex.search(first_cell_text) and len(first_cell_text.split()) < 4:
                        if r_idx == 0 or (r_idx > 0 and table.cell(r_idx,0).text != table.cell(r_idx-1,0).text):
                            current_table_group_label = first_cell_text.replace(":","").strip()
                        elif not current_table_group_label:
                             current_table_group_label = first_cell_text.replace(":","").strip()
                
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for r_match in regex.finditer(cell_text):
                        ph_key = r_match.group(1).strip()
                        if ph_key and ph_key not in schema: # Only process if not already in schema
                            immediate_label = ""
                            if c_idx > 0:
                                label_cell_text = row.cells[c_idx-1].text.strip().replace(":","").strip()
                                if label_cell_text and not regex.search(label_cell_text):
                                    immediate_label = label_cell_text
                            if not immediate_label:
                                immediate_label = cell_text[:r_match.start()].strip().replace(":","").strip()
                            
                            placeholder_details_dynamic[ph_key] = {
                                "immediate_label": immediate_label if immediate_label else ph_key,
                                "section": current_section_header,
                                "subsection": current_subsection_header,
                                "table_group": current_table_group_label if current_table_group_label != immediate_label else "",
                                "type": "table"
                            }
        
        # Merge dynamically found details into the schema
        for ph_key, details in placeholder_details_dynamic.items():
            if ph_key not in schema: # Ensure not to overwrite explicitly mapped ones unless intended
                field_type = "boolean" if ph_key.endswith("_check") else "string"
                description_parts = []
                if details.get("immediate_label") and details["immediate_label"] != ph_key:
                    description_parts.append(details["immediate_label"])
                if details.get("table_group"):
                    description_parts.append(details["table_group"])
                description = " - ".join(filter(None, description_parts)) or ph_key
                
                schema[ph_key] = {
                    "type": field_type,
                    "section": details.get("section", "General"),
                    "subsection": details.get("subsection", ""),
                    "description": description,
                    "location": details.get("type", "unknown")
                }
                if field_type == "boolean":
                    # Use a simplified synonym/indicator generation for dynamically found items
                    synonyms = [description.lower(), ph_key.replace("_check","").replace("_"," ")]
                    schema[ph_key]["synonyms"] = list(set(synonyms))
                    schema[ph_key]["positive_indicators"] = [f"with {s}" for s in synonyms] + ["yes", "selected", description.lower()]

        # Final check: ensure all placeholders found in the document have an entry
        all_phs_from_doc = extract_placeholders(template_path) # Re-extract to get all
        for ph in all_phs_from_doc:
            if ph not in schema:
                field_type = "boolean" if ph.endswith("_check") else "string"
                schema[ph] = {
                    "type": field_type,
                    "section": "General", # Default section
                    "subsection": "",
                    "description": ph, # Default description
                    "location": "unknown" # Default location
                }
                if field_type == "boolean":
                    synonyms = [ph.lower().replace("_check","").replace("_"," ")]
                    schema[ph]["synonyms"] = synonyms
                    schema[ph]["positive_indicators"] = [f"with {s}" for s in synonyms] + ["yes", "selected", ph.lower()]


        print(f"Generated schema for {len(schema)} SortStar placeholders (explicit + dynamic).")
        return schema

    except Exception as e:
        print(f"Error extracting placeholder schema for SortStar (dynamic part): {e}")
        import traceback
        traceback.print_exc()
        # Fallback to schema based only on explicit mappings if dynamic fails
        return {k:v for k,v in schema.items() if k in explicit_placeholder_mappings}


# Functions like generate_synonyms_for_checkbox, generate_positive_indicators, 
# add_section_aware_instructions, enhance_placeholder_context_with_outline, 
# and parse_full_fields_outline might need to be adapted or simplified if the 
# SortStar template and its associated metadata (like an outline file) are different.
# For now, these are omitted or would require SortStar specific versions.
# The primary goal is to get the placeholder filling mechanism working with the new mappings.

if __name__ == '__main__':
    # Path to the SortStar template - this needs to be correctly set for testing
    test_sortstar_template_path = os.path.join("..", "..", "templates", "goa_sortstar_temp.docx") # Adjust path as needed
    
    if os.path.exists(test_sortstar_template_path):
        print(f"Testing with SortStar template: {test_sortstar_template_path}")
        
        placeholders = extract_placeholders(test_sortstar_template_path)
        print("\n--- Unique Placeholders Found (SortStar) ---")
        if placeholders:
            for p_holder in placeholders:
                print(f"- {p_holder}")
        else:
            print("No placeholders found in SortStar template.")

        print("\n--- Placeholder Context (Hierarchical Test - SortStar) ---")
        # For SortStar, we might initially rely more on explicit mappings
        # or a simplified hierarchical extraction if no specific outline exists yet.
        context = extract_placeholder_context_hierarchical(test_sortstar_template_path, enhance_with_outline=False)
        if context:
            count = 0
            for p_holder, desc in sorted(context.items()): 
                print(f"- '{p_holder}': '{desc}'")
                count += 1
                if count >= 50 and len(context) > 50: 
                    print("... and more.")
                    break
        else:
            print("No placeholder context generated for SortStar.")

        print("\n--- Placeholder Schema (SortStar Test) ---")
        schema = extract_placeholder_schema(test_sortstar_template_path)
        if schema:
            count = 0
            for p_holder, sch_info in sorted(schema.items()):
                print(f"- '{p_holder}': Type: {sch_info.get('type')}, Desc: {sch_info.get('description')}")
                count +=1
                if count >= 50 and len(schema) > 50:
                    print("... and more.")
                    break
        else:
            print("No schema generated for SortStar")
            
    else:
        print(f"SortStar test template file not found at: {test_sortstar_template_path}. Please provide a valid path.")
        print(f"Current working directory: {os.getcwd()}")
        # Try an alternative common path if the above fails during testing
        alt_test_template_path = "templates/goa_sortstar_temp.docx"
        if os.path.exists(alt_test_template_path):
            print(f"Attempting with alternative path: {alt_test_template_path}")
            # Rerun test with alternative path (logic duplicated for brevity here)
            # This indicates that the __main__ block might need more robust path handling
            # when run directly vs. when the module is imported.
        else:
            print(f"Alternative path {alt_test_template_path} also not found.") 