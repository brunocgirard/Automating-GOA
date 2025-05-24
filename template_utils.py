import re
from typing import List, Set, Dict, Optional
from docx import Document
import os # Added for __main__ example

def extract_placeholders(template_path: str) -> List[str]:
    """Reads the template and extracts all unique, cleaned placeholder keys."""
    print(f"Extracting placeholders from: {template_path}")
    placeholders: Set[str] = set()
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")

        for para in doc.paragraphs:
            for match in regex.findall(para.text):
                cleaned_key = match.strip()
                if cleaned_key:
                    placeholders.add(cleaned_key)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in regex.findall(cell.text):
                        cleaned_key = match.strip()
                        if cleaned_key:
                            placeholders.add(cleaned_key)
                            
        if not placeholders:
             print(f"Warning: No placeholders found in {template_path}")
             return []

        print(f"Found {len(placeholders)} unique placeholders.")
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error reading placeholders from template '{template_path}': {e}")
        return []

def is_likely_section_header(paragraph) -> bool:
    """Heuristically determines if a paragraph is a section header."""
    text = paragraph.text.strip()
    if not text: # Empty paragraphs are not headers
        return False
    
    # Condition 1: All caps and relatively short
    if text.isupper() and len(text.split()) < 7:
        return True
        
    # Condition 2: Ends with a colon and is relatively short
    # if text.endswith(":") and len(text.split()) < 7:
    #     return True # This might be too broad, catching field labels

    # Condition 3: Bold text (check first run)
    if paragraph.runs and paragraph.runs[0].bold and len(text.split()) < 7:
        # Further check: not too many lowercase letters if it's mostly uppercase (e.g. avoids bolded sentences)
        if sum(1 for char in text if char.islower()) < len(text) / 2:
             return True
    
    # Condition 4: Check for specific heading styles if used consistently in template
    # for style_name_part in ["heading 1", "heading 2", "heading 3", "title"]:
    #     if paragraph.style and paragraph.style.name and style_name_part in paragraph.style.name.lower():
    #         return True
            
    # Add more heuristics as needed (e.g., font size significantly larger than body text)
    return False

def extract_placeholder_context_hierarchical(template_path: str) -> Dict[str, str]:
    """
    Parses the template to extract placeholders and attempts to build hierarchical context
    by identifying section headers.
    """
    print(f"Extracting hierarchical placeholder context from: {template_path}")
    context_map: Dict[str, str] = {}
    try:
        doc = Document(template_path)
        regex = re.compile(r"{{\s*(.*?)\s*}}")
        
        current_section_header = "General" # Default section
        current_subsection_header = ""

        # First pass: Identify all placeholders and their immediate cell/paragraph text
        # Store as: {placeholder: {"immediate": "text", "table_rc": (table_idx, r, c) or None, "para_idx": p_idx or None}}
        placeholder_details = {}

        for p_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if is_likely_section_header(para):
                # Very basic subsection detection: if a new header is found while a section header is active
                # This doesn't handle deep nesting well without style checks.
                # For now, let's assume one level of sectioning for simplicity of header capture.
                # A more robust way would be to check styles (Heading 1, Heading 2, etc.)
                current_section_header = para_text.replace(":","").strip()
                current_subsection_header = "" # Reset subsection on new main section
                # print(f"DEBUG: New Section Header: {current_section_header}")
                continue # Don't look for placeholders in headers themselves
            
            # Heuristic for subsection: If text is bold and short but not ALL CAPS
            elif para.runs and para.runs[0].bold and len(para_text.split()) < 5 and not para_text.isupper() and para_text.endswith(":"):
                current_subsection_header = para_text.replace(":","").strip()
                # print(f"DEBUG: New Sub-Section Header: {current_subsection_header}")
                continue

            for r_match in regex.finditer(para_text):
                ph_key = r_match.group(1).strip()
                if ph_key and ph_key not in placeholder_details:
                    preceding_text = para_text[:r_match.start()].strip()
                    preceding_text = regex.sub("", preceding_text).strip().replace(":","").strip()
                    placeholder_details[ph_key] = {
                        "immediate_label": preceding_text if preceding_text else ph_key,
                        "section": current_section_header,
                        "subsection": current_subsection_header,
                        "type": "paragraph"
                    }

        for t_idx, table in enumerate(doc.tables):
            current_table_group_label = "" # For labels like HMI, PLC in the first column spanning rows
            for r_idx, row in enumerate(table.rows):
                # Try to detect a group label in the first column if it spans or is consistent
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text and not regex.search(first_cell_text) and len(first_cell_text.split()) < 4:
                        # Heuristic: if this cell is different from the one above it in the same column, 
                        # or if it's the first row, it might be a new group label for subsequent rows.
                        if r_idx == 0 or (r_idx > 0 and table.cell(r_idx,0).text != table.cell(r_idx-1,0).text):
                            current_table_group_label = first_cell_text.replace(":","").strip()
                        elif not current_table_group_label: # If still no group label, take from first row, first cell
                            current_table_group_label = first_cell_text.replace(":","").strip()
                
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for r_match in regex.finditer(cell_text):
                        ph_key = r_match.group(1).strip()
                        if ph_key and ph_key not in placeholder_details: # Prioritize paragraph context if already found
                            immediate_label = ""
                            if c_idx > 0: # Label to the left
                                label_cell_text = row.cells[c_idx-1].text.strip().replace(":","").strip()
                                if label_cell_text and not regex.search(label_cell_text):
                                    immediate_label = label_cell_text
                            if not immediate_label: # Text before in current cell
                                immediate_label = cell_text[:r_match.start()].strip().replace(":","").strip()
                            
                            placeholder_details[ph_key] = {
                                "immediate_label": immediate_label if immediate_label else ph_key,
                                "section": current_section_header,
                                "subsection": current_subsection_header, # Could be overridden by table group
                                "table_group": current_table_group_label if current_table_group_label != immediate_label else "",
                                "type": "table"
                            }

        # Construct final context strings
        for ph_key, details in placeholder_details.items():
            parts = []
            if details.get("section") and details["section"] != "General": parts.append(details["section"])
            if details.get("subsection"): parts.append(details["subsection"])
            if details.get("table_group"): parts.append(details["table_group"])
            if details.get("immediate_label") and details["immediate_label"] != ph_key: parts.append(details["immediate_label"])
            
            if parts:
                context_map[ph_key] = " - ".join(filter(None, parts))
            else:
                context_map[ph_key] = ph_key # Fallback

        # Ensure all placeholders from extract_placeholders have some context
        all_phs = extract_placeholders(template_path)
        for ph in all_phs:
            if ph not in context_map:
                context_map[ph] = ph # Default if missed

        if not context_map: print(f"Warning: No placeholder context generated for {template_path}")
        else: print(f"Generated hierarchical context for {len(context_map)} placeholders.")
        return context_map

    except Exception as e:
        print(f"Error extracting hierarchical placeholder context from '{template_path}': {e}")
        import traceback
        traceback.print_exc()
        return {}

if __name__ == '__main__':
    test_template_path = 'template.docx' 
    if os.path.exists(test_template_path):
        print(f"Testing with template: {test_template_path}")
        
        placeholders = extract_placeholders(test_template_path)
        print("\n--- Unique Placeholders Found ---")
        if placeholders:
            for p_holder in placeholders:
                print(f"- {p_holder}")
        else:
            print("No placeholders found.")

        print("\n--- Placeholder Context (Hierarchical Test) ---")
        context = extract_placeholder_context_hierarchical(test_template_path)
        if context:
            count = 0
            for p_holder, desc in sorted(context.items()): 
                print(f"- '{p_holder}': '{desc}'")
                count += 1
                if count >= 50 and len(context) > 50: 
                    print("... and more.")
                    break
        else:
            print("No placeholder context generated.")
            
    else:
        print(f"Test template file not found at: {test_template_path}. Please provide a valid path.") 